"""
Pruebas del orquestador que no cuestan tokens ni tocan nada real.

Cubren la logica determinista (fechas, nombres, flashcards, .docx) y, sobre
todo, las garantias de aislamiento del modo ensayo: que un ensayo no pueda
escribir en el vault real, en Anki, en config.json ni en la carpeta de
intermedios. Esa ultima garantia existe porque ya fallo una vez: un ensayo
dejaba <slug>_skill.json en la carpeta real y eso habria hecho que la
siguiente corrida de verdad se saltara la clase en silencio.

Correr con:
    python3 tests/test_orquestador.py
"""
import shutil
import sys
import tempfile
from datetime import date
from pathlib import Path

RAIZ = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(RAIZ))

from orquestador import carpetas, deteccion, docx_generator, ensayo, extraer_flashcards, nombres
from orquestador.config import PENDIENTES_DIR_POR_DEFECTO, dir_pendientes, usar_dir_pendientes

fallos: list[str] = []


def check(nombre: str, cond: bool, detalle: str = "") -> None:
    print(("  OK   " if cond else "  FALLA") + f"  {nombre}" + (f"  <- {detalle}" if not cond and detalle else ""))
    if not cond:
        fallos.append(nombre)


def probar_nombres() -> None:
    print("\n== nombres ==")
    a = nombres.slug_pendiente("1970-01-20|a.m4a")
    b = nombres.slug_pendiente("1970-01-20|b.m4a")
    check("dos clases con la misma fecha corrupta no colisionan", a != b, f"{a} == {b}")
    check("el slug conserva la fecha legible", a.startswith("1970-01-20"))
    sucio = 'a/b:c*d?e"f<g>h|i'
    check("sanitiza caracteres ilegales de archivo",
          nombres.sanitizar_nombre_archivo(sucio) == "a-b-c-d-e-f-g-h-i")
    check("un titulo con .. no puede escapar de la carpeta",
          "/" not in nombres.sanitizar_nombre_archivo("../../etc/passwd"))
    check("numero de clase con cero a la izquierda",
          nombres.nombre_base(3, "2026-08-05", "Tema") == "Clase 03 - 2026-08-05 - Tema")


def probar_deteccion() -> None:
    print("\n== deteccion de fecha y ramo ==")
    check("rescata la fecha DD.MM.YY del nombre del archivo",
          deteccion._extraer_fecha_de_nombre("CAB7 10.04.25 Performance") == date(2025, 4, 10))
    check("ignora un nombre sin fecha",
          deteccion._extraer_fecha_de_nombre("grabacion final.m4a") is None)
    check("rechaza una fecha imposible", deteccion._extraer_fecha_de_nombre("x 32.13.25 y") is None)

    cfg = {"semestre": {"fecha_inicio": "2026-08-03"},
           "ramos": {"lunes": {"nombre": "R", "perfil_whisper": "es-chile"}}}
    check("una fecha anterior al semestre no asigna ramo (bug de 1970)",
          deteccion.resolver_ramo(date(1970, 1, 20), cfg) is None)
    check("un lunes dentro del semestre si asigna ramo",
          (deteccion.resolver_ramo(date(2026, 8, 3), cfg) or {}).get("nombre") == "R")
    check("un sabado no asigna ramo", deteccion.resolver_ramo(date(2026, 8, 8), cfg) is None)
    check("la semana de semestre se calcula bien",
          deteccion.calcular_semana_semestre(date(2026, 8, 10), "2026-08-03") == 2)


def probar_flashcards() -> None:
    print("\n== extraccion de flashcards ==")
    nota = """## 10 preguntas
(de menor a mayor dificultad, tapate las respuestas)
1. Que es la elasticidad?
2. Por que importa el excedente?

## Respuestas modelo
(las que daria alguien que domina el tema)
1. La sensibilidad de la cantidad ante el precio.
2. Porque mide el bienestar.

## Otra seccion
texto que no debe entrar
"""
    t = extraer_flashcards.extraer_preguntas_respuestas(nota)
    check("extrae exactamente dos tarjetas", len(t) == 2, str(t))
    check("no cuela la linea de instruccion como pregunta",
          all("dificultad" not in p and "domina el tema" not in r for p, r in t))
    check("empareja cada pregunta con su respuesta",
          t[0][0].startswith("Que es la elasticidad") and t[0][1].startswith("La sensibilidad"))
    check("no arrastra la seccion siguiente",
          all("no debe entrar" not in p + r for p, r in t))


def probar_docx() -> None:
    print("\n== generacion del .docx ==")
    md = """---
ramo: PRUEBA
---
# Titulo
Texto con **negrita**.

| Concepto | Definicion |
| --- | --- |
| Elasticidad | Sensibilidad |

## Preguntas de repaso
esto no debe aparecer en el docx
"""
    tmp = Path(tempfile.mkdtemp())
    try:
        ruta = docx_generator.generar_docx(
            {"numero_clase": 1, "fecha": "2026-08-05", "ramo": "PRUEBA"},
            "Clase de prueba",
            "# Fuente\nTexto de respaldo.",
            md,
            [{"concepto": "Elasticidad", "por_que": "se repite al inicio y al cierre"}],
            {"rutas": {"output": str(tmp)}},
        )
        check("crea el archivo .docx", ruta.is_file() and ruta.stat().st_size > 0)
        check("usa el nombre de clase correcto",
              ruta.name == "Clase 01 - 2026-08-05 - Clase de prueba.docx", ruta.name)

        from docx import Document
        doc = Document(str(ruta))
        texto = "\n".join(p.text for p in doc.paragraphs)
        check("renderiza las tablas markdown", len(doc.tables) >= 2, f"tablas={len(doc.tables)}")
        check("omite las secciones de repaso espaciado", "esto no debe aparecer" not in texto)
        check("quita el frontmatter YAML", "ramo: PRUEBA" not in texto)
        check("la negrita queda sin asteriscos", "**negrita**" not in texto and "negrita" in texto)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def probar_aislamiento_del_ensayo() -> None:
    """El bloque que mas importa: que un ensayo no pueda tocar nada real."""
    print("\n== aislamiento del modo ensayo ==")
    config_real = {
        "rutas": {"input": "/real/input", "output": "/real/output",
                  "procesados": "/real/procesados",
                  "vault_obsidian": "/real/vault", "transcriptotem": "/real/tt"},
        "semestre": {"fecha_inicio": "2026-08-03"},
        "ramos": {}, "carpetas_ramo": {"RAMO VIEJO": "/real/vault/RAMO VIEJO"},
    }
    pendientes_antes = sorted(p.name for p in PENDIENTES_DIR_POR_DEFECTO.glob("*")) \
        if PENDIENTES_DIR_POR_DEFECTO.exists() else []

    cfg, sandbox = ensayo.preparar(config_real)
    try:
        for clave in ("input", "output", "procesados", "vault_obsidian"):
            check(f"la ruta '{clave}' apunta al sandbox", str(sandbox) in cfg["rutas"][clave])
        check("el ensayo queda marcado", ensayo.es_ensayo(cfg))
        check("una config normal no se confunde con un ensayo", not ensayo.es_ensayo(config_real))
        check("el cache de carpetas reales se descarta", cfg["carpetas_ramo"] == {})
        check("la config original no se muta",
              config_real["rutas"]["vault_obsidian"] == "/real/vault")
        check("los intermedios se redirigen al sandbox", str(sandbox) in str(dir_pendientes()))

        carpeta, _ = carpetas.resolver_carpeta_ramo("RAMO NUEVO", cfg["rutas"]["vault_obsidian"], cfg)
        check("una carpeta de ramo nueva cae dentro del sandbox", str(sandbox) in str(carpeta))
    finally:
        usar_dir_pendientes(None)
        shutil.rmtree(sandbox, ignore_errors=True)

    check("los intermedios vuelven a la carpeta real al terminar",
          dir_pendientes() == PENDIENTES_DIR_POR_DEFECTO)
    pendientes_despues = sorted(p.name for p in PENDIENTES_DIR_POR_DEFECTO.glob("*")) \
        if PENDIENTES_DIR_POR_DEFECTO.exists() else []
    check("el ensayo no dejo archivos en la carpeta real",
          pendientes_antes == pendientes_despues)


def probar_archivado_no_destructivo() -> None:
    print("\n== el ensayo no mueve el audio original ==")
    from orquestador.archivado import archivar_audio

    tmp = Path(tempfile.mkdtemp())
    try:
        origen = tmp / "clase original.m4a"
        origen.write_bytes(b"audio")
        trabajo = {"ramo": "RAMO", "numero_clase": 1, "fecha": "2026-08-05",
                   "archivos": [str(origen)]}

        cfg_ensayo = {"rutas": {"procesados": str(tmp / "dest")}, ensayo.CLAVE: True}
        archivar_audio(trabajo, "Titulo", cfg_ensayo)
        check("en ensayo el audio original sigue en su lugar", origen.is_file())
        check("en ensayo si queda una copia archivada",
              len(list((tmp / "dest").rglob("*.m4a"))) == 1)

        cfg_real = {"rutas": {"procesados": str(tmp / "dest_real")}}
        archivar_audio(trabajo, "Titulo", cfg_real)
        check("fuera de ensayo el audio si se mueve", not origen.exists())
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def probar_dialogo_nunca_descarta_solo() -> None:
    """
    Descartar una grabacion tiene que costar un clic explicito. Esto ya fallo
    una vez con una clase real: el timeout y el boton por defecto llevaban los
    dos a ignorar, y el audio se archivo solo mientras nadie miraba.
    """
    print("\n== el dialogo nunca descarta por su cuenta ==")
    from orquestador import dialogo_no_reconocido as dlg

    original = dlg._mostrar_dialogo_principal
    casos = [
        ("timeout o ventana cerrada", None, "solo_transcribir"),
        ("clic explicito en Ignorar", dlg.OPCION_IGNORAR, "ignorar"),
        ("clic en Solo transcribir", dlg.OPCION_SOLO_TRANSCRIBIR, "solo_transcribir"),
    ]
    try:
        for nombre, respuesta, esperado in casos:
            dlg._mostrar_dialogo_principal = lambda _t, _r=respuesta: _r
            accion = dlg.preguntar_que_hacer({"archivos": [], "fecha": "1970-01-20",
                                              "dia_semana": "martes"}, {})["accion"]
            check(f"{nombre} -> {esperado}", accion == esperado, f"dio '{accion}'")

        # Abandonar a mitad de elegir el ramo tampoco puede descartar.
        dlg._mostrar_dialogo_principal = lambda _t: dlg.OPCION_APLICAR_SKILLS
        original_ramo = dlg._elegir_ramo
        dlg._elegir_ramo = lambda _c: None
        try:
            accion = dlg.preguntar_que_hacer({"archivos": [], "fecha": "1970-01-20",
                                              "dia_semana": "martes"}, {})["accion"]
            check("abandonar la eleccion de ramo -> solo_transcribir",
                  accion == "solo_transcribir", f"dio '{accion}'")
        finally:
            dlg._elegir_ramo = original_ramo

    finally:
        dlg._mostrar_dialogo_principal = original

    import inspect
    fuente = inspect.getsource(original)
    check("el boton por defecto del dialogo no es el que descarta",
          "default button {_escapar(OPCION_IGNORAR)}" not in fuente
          and "OPCION_SOLO_TRANSCRIBIR" in fuente)


def probar_el_nombre_del_archivo_manda_sobre_el_dia() -> None:
    """
    El caso real que esto arregla: el 26-08-2026 (miercoles) se subieron dos
    grabaciones, la clase de DESEMPENO ORGANIZACIONAL y una reunion informativa
    de un ramo anexo. Como el dia tenia ramo asignado, las dos se archivaron
    como esa clase, sin preguntar, y chocaron de numero.

    Los nombres de archivo de aqui son los reales que pasaron por el sistema,
    con las faltas de ortografia incluidas. Son la calibracion del umbral: si
    alguien lo mueve, esto tiene que seguir pasando.
    """
    print("\n== el nombre del archivo manda sobre el dia de la semana ==")
    from orquestador import ramo_por_nombre as rpn

    config = {
        "ramos": {
            "lunes": {"nombre": "MARKETING ESTRATÉGICO", "perfil_whisper": "es-chile"},
            "martes": {"nombre": "MERCADOS Y ESTRUCTURA ECONÓMICA", "perfil_whisper": "es-chile"},
            "miercoles": {"nombre": "DESEMPEÑO ORGANIZACIONAL", "perfil_whisper": "es-chile"},
            "jueves": {"nombre": "ECONOMETRÍA", "perfil_whisper": "es-chile"},
            "viernes": {"nombre": "TALLER EN BUSINESS ANALYTICS III", "perfil_whisper": "es-chile"},
        },
        "semestre": {"fecha_inicio": "2026-08-03"},
    }

    reconoce = [
        ("Marketing estratégico 24.08.26.m4a", "MARKETING ESTRATÉGICO"),
        ("Econometría 13.08.26.m4a", "ECONOMETRÍA"),
        ("Desempeño organizacional 26.08.26.m4a", "DESEMPEÑO ORGANIZACIONAL"),
        # Con la falta de ortografia del estudiante y sin la palabra "Taller".
        ("Business Analitics III 21.08.26.m4a", "TALLER EN BUSINESS ANALYTICS III"),
        # "Mercado" en singular.
        ("Mercado y estructura económica 18.8.26.m4a", "MERCADOS Y ESTRUCTURA ECONÓMICA"),
        # Otros formatos de fecha que el estudiante usa.
        ("Marketing estratégico 3-8-26.m4a", "MARKETING ESTRATÉGICO"),
        # Sufijo de parte que agrega el propio sistema al cortar un audio largo.
        ("Marketing estratégico 10.08.26 - parte 03.m4a", "MARKETING ESTRATÉGICO"),
    ]
    for nombre, esperado in reconoce:
        estado, info = rpn.resolver([nombre], config)
        check(f"'{nombre[:34]}' -> {esperado[:24]}",
              estado == rpn.RECONOCIDO and info["nombre"] == esperado,
              f"dio {estado} / {info}")

    # El caso que motivo todo: palabras reales que no son ningun ramo del
    # horario. No puede caer al dia de la semana.
    pregunta = [
        "Medición competencias intermedias 26.08.26.m4a",
        "Reunión secretario académico 30.08.m4a",
        "Conversación con Javier.m4a",
    ]
    for nombre in pregunta:
        estado, info = rpn.resolver([nombre], config)
        check(f"'{nombre[:34]}' -> pregunta",
              estado == rpn.NO_CALZA and info is None, f"dio {estado} / {info}")

    # Un nombre sin contenido no es evidencia de nada: ahi si manda el dia.
    for nombre in ["Nota de voz 3.m4a", "Grabación 12.m4a", "audio.m4a", "Clase.m4a"]:
        estado, _ = rpn.resolver([nombre], config)
        check(f"'{nombre}' -> sin senal, decide el dia", estado == rpn.SIN_SENAL,
              f"dio {estado}")

    # Un ramo agregado a mano tiene que reconocerse igual que los del horario.
    con_adicional = dict(config)
    con_adicional["ramos_adicionales"] = {
        "MIC - MEDICIÓN INTERMEDIA DE COMPETENCIAS": {"perfil_whisper": "es-chile"}
    }
    estado, info = rpn.resolver(["Medición competencias intermedias 26.08.26.m4a"], con_adicional)
    check("una vez creado el ramo, la siguiente reunion se reconoce sola",
          estado == rpn.RECONOCIDO and info["nombre"].startswith("MIC"),
          f"dio {estado} / {info}")

    # Partes de ramos distintos agrupadas: ya fusiono dos clases una vez.
    estado, _ = rpn.resolver(
        ["Desempeño organizacional 19.08.26.m4a", "Econometría 20.08.26.m4a"], config)
    check("dos ramos distintos en un mismo trabajo -> pregunta", estado == rpn.NO_CALZA,
          f"dio {estado}")

    # Y el efecto completo sobre la deteccion: el mismo miercoles, dos
    # grabaciones, cada una a su lugar.
    miercoles = date(2026, 8, 26)
    clase = deteccion.resolver_ramo_de_grabacion(
        ["Desempeño organizacional 26.08.26.m4a"], miercoles, config)
    reunion = deteccion.resolver_ramo_de_grabacion(
        ["Medición competencias intermedias 26.08.26.m4a"], miercoles, config)
    check("el miercoles, la clase real sigue siendo la clase",
          clase and clase["nombre"] == "DESEMPEÑO ORGANIZACIONAL", f"dio {clase}")
    check("el mismo miercoles, la reunion no se archiva sola",
          reunion is None, f"dio {reunion}")

    # Sin nombre util, el comportamiento de siempre: manda el dia.
    por_dia = deteccion.resolver_ramo_de_grabacion(["Nota de voz 3.m4a"], miercoles, config)
    check("sin nombre util se mantiene el comportamiento de siempre",
          por_dia and por_dia["nombre"] == "DESEMPEÑO ORGANIZACIONAL", f"dio {por_dia}")


def probar_bitacora_deshace_todo() -> None:
    """
    El aborto promete dejar el disco como estaba. Si el deshacer falla, esa
    promesa se rompe justo cuando el estudiante ya decidio cancelar, o sea en
    el peor momento para descubrirlo.
    """
    print("\n== abortar deja todo como estaba ==")
    from orquestador.bitacora import Bitacora

    base = Path(tempfile.mkdtemp())
    try:
        # Estado inicial: una carpeta de ramo con una nota que ya existia.
        vault = base / "vault" / "RAMO"
        vault.mkdir(parents=True)
        indice = vault / "indice.md"
        indice.write_text("contenido original", encoding="utf-8")

        entrada = base / "Input"
        entrada.mkdir()
        audio = entrada / "clase.m4a"
        audio.write_bytes(b"audio")

        b = Bitacora(base / "bitacora.json")

        # Lo que hace una corrida.
        b.fotografiar_carpeta(vault)
        (vault / "Fuente - nueva.md").write_text("nota nueva", encoding="utf-8")
        (vault / "Aprendizaje - nueva.md").write_text("otra nota", encoding="utf-8")
        indice.write_text("contenido MODIFICADO por la skill", encoding="utf-8")

        carpeta_nueva = base / "vault" / "RAMO RECIEN CREADO"
        carpeta_nueva.mkdir()
        b.carpeta_creada(carpeta_nueva)

        docx = base / "salida.docx"
        docx.write_bytes(b"docx")
        b.archivo_creado(docx)

        destino = base / "Procesados" / "clase archivada.m4a"
        destino.parent.mkdir(parents=True)
        b.audio_movido(audio, destino)
        shutil.move(str(audio), str(destino))

        check("durante la corrida el audio ya no esta en Input", not audio.exists())

        # Abortar.
        revertido = b.deshacer()

        check("las notas nuevas se borraron",
              not (vault / "Fuente - nueva.md").exists()
              and not (vault / "Aprendizaje - nueva.md").exists())
        check("la nota que ya existia sigue ahi", indice.is_file())
        check("y con su contenido original",
              indice.read_text(encoding="utf-8") == "contenido original",
              indice.read_text(encoding="utf-8"))
        check("la carpeta creada se elimino", not carpeta_nueva.exists())
        check("el .docx se borro", not docx.exists())
        check("el audio volvio a Input", audio.is_file() and not destino.exists())
        check("el audio conserva su contenido", audio.read_bytes() == b"audio")
        check("se informa lo que se revirtio", len(revertido) >= 4, str(revertido))
        check("la bitacora se limpia al terminar", not (base / "bitacora.json").exists())
    finally:
        shutil.rmtree(base, ignore_errors=True)


def probar_bitacora_no_borra_lo_ajeno() -> None:
    print("\n== abortar no toca lo que no puso el sistema ==")
    from orquestador.bitacora import Bitacora

    base = Path(tempfile.mkdtemp())
    try:
        carpeta = base / "RAMO"
        carpeta.mkdir()
        b = Bitacora(base / "b.json")
        b.carpeta_creada(carpeta)
        ajeno = carpeta / "algo mio.md"
        ajeno.write_text("no es del sistema", encoding="utf-8")

        b.deshacer()
        check("una carpeta con algo dentro no se borra", carpeta.is_dir())
        check("el archivo ajeno sigue intacto", ajeno.is_file())
    finally:
        shutil.rmtree(base, ignore_errors=True)


def probar_seccion_critica() -> None:
    """Durante el movimiento del audio el aborto se encola, no se ignora."""
    print("\n== el tramo delicado no se puede cortar por la mitad ==")
    from orquestador import cancelacion

    llego_al_final = False
    try:
        with cancelacion.seccion_critica():
            check("mientras dura, el estado dice que no se puede interrumpir",
                  cancelacion._seccion_critica.locked())
            # Simula el pedido de aborto llegando justo aqui.
            cancelacion._al_recibir_senal(15, None)
            llego_al_final = True
    except cancelacion.Abortado:
        check("el aborto pedido adentro se aplica al salir, no se pierde", True)
    else:
        check("el aborto pedido adentro se aplica al salir, no se pierde", False)

    check("el tramo se completo antes de abortar", llego_al_final)
    check("el candado queda libre despues", not cancelacion._seccion_critica.locked())


def probar_formulas() -> None:
    print("\n== formulas con subindices y superindices reales ==")
    from docx import Document as Doc
    from orquestador import formulas

    check("reconoce una formula destacada", formulas.es_formula_destacada("$$x̄ ± σ$$"))
    check("no confunde texto normal con formula", not formulas.es_formula_destacada("cuesta $5 o $8"))
    check("extrae el contenido", formulas.texto_de_formula_destacada("$$x̄$$") == "x̄")

    tmp = Path(tempfile.mkdtemp())
    try:
        # Una formula destacada se dibuja como imagen, con tipografia
        # matematica de verdad (barra de fraccion, radical que se estira).
        doc = Doc()
        formulas.agregar_formula_destacada(
            doc, r"\bar{x} \pm z_{1-\alpha/2} \cdot \frac{\sigma}{\sqrt{n}}")
        ruta = Path(tmp) / "conformula.docx"
        doc.save(str(ruta))
        check("la formula destacada queda como imagen",
              len(Doc(str(ruta)).inline_shapes) == 1,
              f"imagenes={len(Doc(str(ruta)).inline_shapes)}")

        # Si el dibujo falla, no se pierde la formula: se escribe como texto.
        original = formulas._renderizar_imagen
        formulas._renderizar_imagen = lambda _l: None
        try:
            doc_fb = Doc()
            formulas.agregar_formula_destacada(doc_fb, "s^{2} = Σ(x_i - x̄)^{2}")
            p = doc_fb.paragraphs[-1]
            check("sin imagen se cae a texto con superindices",
                  [r.text for r in p.runs if r.font.superscript] == ["2", "2"])
            check("sin imagen se cae a texto con subindices",
                  [r.text for r in p.runs if r.font.subscript] == ["i"])
            check("y conserva los simbolos Unicode", "Σ" in p.text and "x̄" in p.text)
        finally:
            formulas._renderizar_imagen = original

        # Una formula corta dentro de una frase sigue siendo texto: una imagen
        # ahi quedaria desalineada con el renglon.
        doc2 = Doc()
        par = doc2.add_paragraph()
        from orquestador.docx_generator import _agregar_texto_con_negritas
        _agregar_texto_con_negritas(par, "La varianza $s^{2}$ usa **n-1** abajo.")
        check("la formula inline se formatea",
              [r.text for r in par.runs if r.font.superscript] == ["2"])
        check("la negrita sigue funcionando junto a la formula",
              any(r.bold and r.text == "n-1" for r in par.runs))
        check("un guion bajo suelto fuera de $ no se toca",
              "_" in _texto_render(doc2, "archivo_de_prueba.txt"))
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def _texto_render(doc, texto: str) -> str:
    from orquestador.docx_generator import _agregar_texto_con_negritas
    p = doc.add_paragraph()
    _agregar_texto_con_negritas(p, texto)
    return p.text


def probar_contexto_va_primero() -> None:
    """El contexto solo sirve si se lee antes de la clase, o sea si esta al
    principio del documento."""
    print("\n== la seccion de contexto va al frente ==")
    from docx import Document as Doc

    tmp = Path(tempfile.mkdtemp())
    try:
        ruta = docx_generator.generar_docx(
            {"numero_clase": 1, "fecha": "2026-08-05", "ramo": "R"},
            "Titulo", "# Fuente\ntexto", "# Aprendizaje\ntexto",
            [{"concepto": "C", "por_que": "p"}],
            {"rutas": {"output": str(tmp)}},
            "# Contexto previo\n\nAlgo que hay que saber antes.",
        )
        encabezados = [p.text for p in Doc(str(ruta)).paragraphs
                       if p.style.name.startswith("Heading") and p.text.strip()]
        check("aparece la seccion de contexto",
              any("Antes de empezar" in h for h in encabezados), str(encabezados))
        check("va antes que los conceptos repetidos",
              encabezados.index("Antes de empezar: lo que conviene tener claro")
              < next(i for i, h in enumerate(encabezados) if "repetidos" in h))

        # Sin contexto el documento se arma igual: la seccion es opcional.
        ruta2 = docx_generator.generar_docx(
            {"numero_clase": 2, "fecha": "2026-08-06", "ramo": "R"},
            "Otro", "# Fuente\ntexto", "# Aprendizaje\ntexto", [],
            {"rutas": {"output": str(tmp)}}, "",
        )
        sin = [p.text for p in Doc(str(ruta2)).paragraphs if p.style.name.startswith("Heading")]
        check("sin contexto el .docx se arma igual",
              ruta2.is_file() and not any("Antes de empezar" in h for h in sin))
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def probar_formulas_en_tablas() -> None:
    """
    Las celdas se escribian como texto plano, asi que la tabla de "que formula
    uso en cada caso" era el unico lugar del documento donde las formulas se
    veian crudas, con el guion bajo y los parentesis a la vista.
    """
    print("\n== formulas dentro de tablas ==")
    from docx import Document as Doc
    from orquestador import formulas

    check("una celda que es toda formula se reconoce",
          formulas.parece_solo_formula("x̄ ± z_(1-α/2)·σ/√n"))
    check("una celda de texto normal no se confunde con formula",
          not formulas.parece_solo_formula(
              "Quiero estimar μ y conozco σ (dato del enunciado)"))

    # El modelo no siempre escribe LaTeX: hay que entender su Unicode.
    latex = formulas.a_latex("p̂ ± z_(1-α/2)·√(p̂q̂/n)")
    check("el sombrero Unicode pasa a LaTeX", r"\hat{p}" in latex, latex)
    check("la raiz Unicode pasa a LaTeX", r"\sqrt{" in latex, latex)
    check("el subindice con parentesis pasa a llaves", "_{1-" in latex, latex)
    check("lo que ya viene en LaTeX no se toca",
          formulas.a_latex(r"\bar{x} \pm \sigma") == r"\bar{x} \pm \sigma")

    md = (
        "| Situación | Fórmula |\n"
        "|---|---|\n"
        "| Conozco σ, muestra grande | x̄ ± z_(1-α/2)·σ/√n |\n"
    )
    tmp = Path(tempfile.mkdtemp())
    try:
        ruta = docx_generator.generar_docx(
            {"numero_clase": 1, "fecha": "2026-08-05", "ramo": "R"},
            "T", "", "# A\n\n" + md, [], {"rutas": {"output": str(tmp)}}, "",
        )
        doc = Doc(str(ruta))
        marca = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"
        imagenes_en_celdas = sum(
            len(p._p.findall(f".//{marca}"))
            for t in doc.tables for f in t.rows for c in f.cells for p in c.paragraphs
        )
        check("la formula de la celda se dibuja", imagenes_en_celdas == 1,
              f"imagenes={imagenes_en_celdas}")
        textos = [c.text for t in doc.tables for f in t.rows for c in f.cells]
        check("ya no queda la formula cruda en el texto",
              not any("z_(1-" in x for x in textos), str(textos))
        check("la celda de texto normal sigue siendo texto",
              any("Conozco σ" in x for x in textos), str(textos))
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def probar_matematica_dentro_de_frases() -> None:
    """
    El modelo mete formulas en medio de las frases sin delimitarlas. Pedirle
    que las marque ya fallo, asi que se reconocen solas, pero sin arrastrar
    palabras normales que llevan guion bajo.
    """
    print("\n== subindices en medio de una frase ==")
    from docx import Document as Doc
    from orquestador.docx_generator import _agregar_texto_con_negritas

    doc = Doc()
    p = doc.add_paragraph()
    _agregar_texto_con_negritas(
        p, "el valor crítico que se busca en tabla (z_(1-α/2) o t_(n-1, 1-α/2))")
    subs = [r.text for r in p.runs if r.font.subscript]
    check("los dos subindices se aplican", subs == ["1-α/2", "n-1, 1-α/2"], str(subs))
    check("ya no queda la notacion cruda", "_(" not in p.text, p.text)

    # Lo que NO debe tocarse.
    for texto in ("El archivo_de_prueba.txt quedó guardado",
                  "usa snake_case para nombrar variables"):
        p2 = doc.add_paragraph()
        _agregar_texto_con_negritas(p2, texto)
        check(f"no toca '{texto[:22]}...'",
              not [r for r in p2.runs if r.font.subscript] and p2.text == texto)

    # La negrita convive con la matematica.
    p3 = doc.add_paragraph()
    _agregar_texto_con_negritas(p3, "usa **z_(α/2)** para el margen")
    check("negrita y subindice a la vez",
          any(r.bold and r.font.subscript for r in p3.runs))


def probar_mapa_y_secciones() -> None:
    """El mapa se dibuja de verdad, y la sesion por bloques de tiempo cede su
    lugar en el .docx a la materia ya digerida."""
    print("\n== mapa dibujado y secciones del .docx ==")
    from docx import Document as Doc

    aprendizaje = (
        "# Aprendizaje\n\n"
        "## 5. Materia lista para estudiar\nExplicacion con ejemplos.\n\n"
        "## Sesión de estudio de 90 minutos\nESTO_NO_VA_AL_DOCX\n\n"
        "## Mapa visual\n"
        "```mapa\n"
        '{"centro": "Tema", "ramas": [{"titulo": "Rama A", "puntos": ["p1", "p2"]},'
        ' {"titulo": "Rama B", "puntos": ["p3"]}]}\n'
        "```\n"
    )
    tmp = Path(tempfile.mkdtemp())
    try:
        ruta = docx_generator.generar_docx(
            {"numero_clase": 1, "fecha": "2026-08-05", "ramo": "R"},
            "T", "# Fuente\ntexto", aprendizaje, [], {"rutas": {"output": str(tmp)}}, "",
        )
        doc = Doc(str(ruta))
        texto = "\n".join(p.text for p in doc.paragraphs)
        encabezados = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]

        check("la sesion por bloques no va al .docx", "ESTO_NO_VA_AL_DOCX" not in texto)
        check("la materia lista para estudiar si va",
              any("Materia lista para estudiar" in h for h in encabezados))
        check("el mapa se inserta como imagen", len(doc.inline_shapes) == 1,
              f"imagenes={len(doc.inline_shapes)}")
        check("no se duplica el titulo del mapa",
              sum(1 for h in encabezados if "mapa" in h.lower()) == 1,
              str([h for h in encabezados if "mapa" in h.lower()]))
        check("el bloque de datos crudo no se imprime", '"centro"' not in texto)

        # Datos mal formados: el documento se arma igual, sin mapa.
        malo = "# A\n\n## Mapa visual\n```mapa\nesto no es json\n```\n"
        r2 = docx_generator.generar_docx(
            {"numero_clase": 2, "fecha": "2026-08-06", "ramo": "R"},
            "T2", "", malo, [], {"rutas": {"output": str(tmp)}}, "",
        )
        check("un mapa mal formado no rompe el documento",
              r2.is_file() and len(Doc(str(r2)).inline_shapes) == 0)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def probar_llamados_a_la_accion() -> None:
    """
    La primera seccion del documento: lo que el profesor pidio.

    Es la parte con mas riesgo del sistema, porque es la que el estudiante mas
    va a creer y la que decide que estudia. Por eso se prueba que la cita
    textual siempre viaje con el punto, que lo dudoso se avise, y que la
    seccion aparezca incluso vacia: "no hubo anuncios" y "el sistema no los
    detecto" no se pueden confundir.
    """
    print("\n== lo que el profesor pidio, al inicio del documento ==")
    from docx import Document as Doc

    llamados = {
        "avisos": [
            {"que": "AVISO_TRABAJO", "cuando": "la proxima semana",
             "textual": "CITA_DEL_TRABAJO", "seguro": True},
            {"que": "AVISO_DUDOSO", "cuando": "", "textual": "CITA_DUDOSA",
             "seguro": False},
        ],
        "evaluacion": [
            {"tema": "TEMA_QUE_ENTRA", "textual": "CITA_DE_LA_PRUEBA", "seguro": True},
        ],
    }

    tmp = Path(tempfile.mkdtemp())
    try:
        ruta = docx_generator.generar_docx(
            {"numero_clase": 1, "fecha": "2026-08-05", "ramo": "R"},
            "T", "# F\ntexto", "# A\ntexto", [{"concepto": "C", "por_que": "p"}],
            {"rutas": {"output": str(tmp)}}, "# Contexto\nprevio", llamados,
        )
        doc = Doc(str(ruta))
        encabezados = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        texto = "\n".join(p.text for p in doc.paragraphs)

        check("aparece la seccion", any("profesor pidió" in h for h in encabezados))
        check("va antes que todo lo demas",
              next(i for i, h in enumerate(encabezados) if "profesor pidió" in h)
              < next(i for i, h in enumerate(encabezados) if "Antes de empezar" in h))
        check("lo que entra en evaluacion tiene su propia subseccion",
              any("entra en evaluación" in h for h in encabezados))

        check("el aviso esta", "AVISO_TRABAJO" in texto)
        check("con su fecha, como la dijo el profe", "la proxima semana" in texto)
        check("y con la cita textual", "CITA_DEL_TRABAJO" in texto)
        check("lo que entra en la prueba esta", "TEMA_QUE_ENTRA" in texto)
        check("con su cita", "CITA_DE_LA_PRUEBA" in texto)
        check("lo dudoso viene avisado",
              "AVISO_DUDOSO" in texto and "no permite estar seguro" in texto)

        # Sin anuncios la seccion sigue estando, con una linea que lo dice.
        r2 = docx_generator.generar_docx(
            {"numero_clase": 2, "fecha": "2026-08-06", "ramo": "R"},
            "T2", "", "# A\ntexto", [], {"rutas": {"output": str(tmp)}}, "", None,
        )
        doc2 = Doc(str(r2))
        texto2 = "\n".join(p.text for p in doc2.paragraphs)
        check("sin anuncios la seccion no desaparece",
              any("profesor pidió" in p.text for p in doc2.paragraphs
                  if p.style.name.startswith("Heading")))
        check("y dice explicitamente que no hubo", "no anunció" in texto2)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def probar_documento_no_se_repite() -> None:
    """
    El documento llego a tener 33 paginas contando la misma materia cuatro
    veces. Contar lo mismo varias veces no refuerza: desplaza a lo que si
    rinde (ver references/diseno-documento.md). Esto comprueba los tres cortes
    que lo dejaron en un tercio.
    """
    print("\n== el documento no cuenta lo mismo dos veces ==")
    from docx import Document as Doc

    # Nota con la estructura vieja: los pasos 1, 2 y 5 por separado.
    aprendizaje = (
        "---\nramo: R\n---\n\n"
        "# Aprendizaje - Titulo Largo Del Archivo\n\n"
        "## 1. Conceptos centrales\nCONCEPTOS_REPETIDOS_TRES_VECES\n\n"
        "## 2. Que dominar para enseñarlo desde cero\nOTRA_VEZ_LO_MISMO\n\n"
        "## 3. Diez preguntas para ponerme a prueba\nLAS_PREGUNTAS\n\n"
        "## 4. Respuestas modelo\nLAS_RESPUESTAS\n\n"
        "## 5. Materia lista para estudiar\nLA_MATERIA_DESARROLLADA\n\n"
        "## 6. Kit de repaso\nEL_KIT\n"
    )
    fuente = (
        "# Fuente - Titulo Largo Del Archivo\n\n"
        "## Resumen\nRESUMEN_CRONOLOGICO\n\n"
        "## Desarrollo\nLA_CLASE_OTRA_VEZ_EN_ORDEN\n\n"
        "## Definiciones y ejemplos del profe\nLAS_DEFINICIONES\n\n"
        "## Huecos y dudas\nLOS_HUECOS\n"
    )

    tmp = Path(tempfile.mkdtemp())
    try:
        ruta = docx_generator.generar_docx(
            {"numero_clase": 1, "fecha": "2026-08-05", "ramo": "R"},
            "T", fuente, aprendizaje, [], {"rutas": {"output": str(tmp)}}, "", None,
        )
        doc = Doc(str(ruta))
        texto = "\n".join(p.text for p in doc.paragraphs)
        encabezados = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]

        check("la materia desarrollada se queda", "LA_MATERIA_DESARROLLADA" in texto)
        check("las dos versiones que la repetian se van",
              "CONCEPTOS_REPETIDOS_TRES_VECES" not in texto
              and "OTRA_VEZ_LO_MISMO" not in texto)
        check("el desarrollo cronologico se va",
              "LA_CLASE_OTRA_VEZ_EN_ORDEN" not in texto
              and "RESUMEN_CRONOLOGICO" not in texto)
        check("las definiciones y los huecos se quedan",
              "LAS_DEFINICIONES" in texto and "LOS_HUECOS" in texto)

        check("primero la materia y despues las preguntas",
              texto.index("LA_MATERIA_DESARROLLADA") < texto.index("LAS_PREGUNTAS"))
        check("y las respuestas despues de las preguntas",
              texto.index("LAS_PREGUNTAS") < texto.index("LAS_RESPUESTAS"))

        check("los encabezados pierden la numeracion del metodo",
              not any(h.strip().startswith(("1.", "2.", "3.", "4.", "5.", "6."))
                      for h in encabezados))
        check("el nombre del archivo de Obsidian no queda de titulo",
              "Titulo Largo Del Archivo" not in texto)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def probar_canales_de_enfasis() -> None:
    """
    Solo hay dos marcas de enfasis y significan cosas distintas. Destacar
    funciona porque es escaso: si todo resalta, no resalta nada.
    """
    print("\n== los dos canales de enfasis ==")
    from docx import Document as Doc

    nota = (
        "# A\n\n## La materia\n"
        "> [!examen] ESTO_ENTRA_EN_LA_PRUEBA\n\n"
        "texto normal\n\n"
        "> [!verificar] ESTO_PUEDE_ESTAR_MAL\n"
    )
    tmp = Path(tempfile.mkdtemp())
    try:
        ruta = docx_generator.generar_docx(
            {"numero_clase": 1, "fecha": "2026-08-05", "ramo": "R"},
            "T", "", nota, [], {"rutas": {"output": str(tmp)}}, "", None,
        )
        doc = Doc(str(ruta))
        texto = "\n".join(p.text for p in doc.paragraphs)
        check("lo que entra en la prueba se etiqueta como tal",
              "ENTRA EN LA PRUEBA" in texto and "ESTO_ENTRA_EN_LA_PRUEBA" in texto)
        check("el aviso de confiabilidad usa la otra marca",
              "Verificar" in texto and "ESTO_PUEDE_ESTAR_MAL" in texto)
        check("el aviso va donde esta el problema, no al principio",
              texto.index("ESTO_ENTRA_EN_LA_PRUEBA") < texto.index("ESTO_PUEDE_ESTAR_MAL"))
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def probar_parrafos_y_justificado() -> None:
    """
    Las notas vienen con el texto cortado cada ochenta o noventa caracteres,
    que es lo normal en markdown. Cada una de esas lineas se convertia en un
    parrafo suelto de Word con su espacio debajo, asi que un parrafo salia
    partido en seis trozos cortados a mitad de frase.
    """
    print("\n== los parrafos no se parten por renglon ==")
    from docx import Document as Doc
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    nota = (
        "# A\n\n## La materia\n"
        "La pregunta que abre la clase es incomoda a proposito: si una empresa\n"
        "tiene buenos productos y buenos servicios, por que igual puede quebrar?\n"
        "La respuesta es que el problema casi nunca es el producto.\n\n"
        "Un segundo parrafo, aparte del primero.\n\n"
        "- una viñeta\n"
        "- otra viñeta\n"
    )
    tmp = Path(tempfile.mkdtemp())
    try:
        ruta = docx_generator.generar_docx(
            {"numero_clase": 1, "fecha": "2026-08-05", "ramo": "R"},
            "T", "", nota, [], {"rutas": {"output": str(tmp)}}, "", None,
        )
        doc = Doc(str(ruta))
        cuerpo = [p.text for p in doc.paragraphs
                  if p.style.name == "Normal" and p.text.strip()]

        check("el parrafo queda entero en un solo parrafo",
              any("incomoda a proposito" in t and "casi nunca es el producto" in t
                  for t in cuerpo))
        check("la linea en blanco si separa parrafos",
              any(t.strip() == "Un segundo parrafo, aparte del primero." for t in cuerpo))
        check("las viñetas no se fusionan entre si",
              sum(1 for p in doc.paragraphs if p.style.name == "List Bullet") == 2)

        # Una viñeta larga tambien viene cortada en la nota, y su segunda mitad
        # se desprendia como parrafo suelto. Se veia igual que el problema que
        # esta funcion venia a arreglar (37 viñetas asi en una nota real).
        from orquestador.docx_generator import _unir_lineas_de_parrafo
        r = _unir_lineas_de_parrafo(
            "- **Evaluacion:** el 30% son talleres\n  y el 30% es el examen.\n"
            "- Otra viñeta.\n"
        ).splitlines()
        check("una viñeta cortada se rearma entera",
              any(l.startswith("- **Evaluacion:**") and "es el examen" in l for l in r))
        check("y no deja huerfano el resto",
              not any(l.startswith("y el 30%") for l in r))
        check("la viñeta siguiente sigue siendo otra", "- Otra viñeta." in r)

        r2 = _unir_lineas_de_parrafo("1. Primer paso que sigue\n   abajo.\n2. Segundo.\n")
        check("los puntos numerados tambien se rearman",
              any(l.startswith("1.") and "abajo." in l for l in r2.splitlines()))

        r3 = _unir_lineas_de_parrafo("> [!examen] entra esto\n> y esto tambien.\n")
        check("una cita de dos renglones es una sola cita",
              len([l for l in r3.splitlines() if l.strip()]) == 1)

        r4 = _unir_lineas_de_parrafo("| a | b |\n|---|---|\n| 1 | 2 |\n")
        check("las tablas no se pegan entre si",
              len([l for l in r4.splitlines() if l.strip()]) == 3)

        r5 = _unir_lineas_de_parrafo('```mapa\n{"centro": "x",\n "ramas": []}\n```\n')
        check("dentro de un bloque cercado no se toca nada",
              '{"centro": "x",' in r5.splitlines())

        alineacion = doc.styles["Normal"].paragraph_format.alignment
        check("el cuerpo va justificado", alineacion == WD_ALIGN_PARAGRAPH.JUSTIFY)

        xml = doc.settings.element.xml
        check("con particion de palabras, para que no queden rios de espacios",
              "autoHyphenation" in xml)

        margen = doc.sections[0].left_margin.cm
        check("y con la linea a un ancho legible", 3.0 <= margen <= 4.0, f"{margen}cm")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def probar_audio_largo_se_corta_solo() -> None:
    """
    Una clase real de 1 h 56 min hizo que Whisper devolviera cero caracteres sin
    lanzar ningun error, y el pipeline siguio hasta entregarle un archivo vacio
    al modelo. Ahora los audios largos se cortan antes de transcribir, y una
    transcripcion vacia detiene la corrida en vez de gastar una llamada.
    """
    print("\n== audios largos y transcripciones vacias ==")
    import subprocess as sp

    from orquestador import transcripcion as tr
    from orquestador.estado_vivo import duracion_audio_segundos

    check("el umbral deja pasar una clase corta sin cortarla",
          tr.UMBRAL_PARTIR_SEGUNDOS > 30 * 60)
    check("el trozo no supera lo ya probado", tr.DURACION_TROZO_SEGUNDOS <= 20 * 60)

    tmp = Path(tempfile.mkdtemp())
    try:
        audio = tmp / "largo.m4a"
        sp.run(["ffmpeg", "-v", "error", "-f", "lavfi",
                "-i", "sine=frequency=440:duration=50",
                "-c:a", "aac", str(audio), "-y"], check=True, capture_output=True)

        original = tr.DURACION_TROZO_SEGUNDOS
        tr.DURACION_TROZO_SEGUNDOS = 20
        try:
            trozos = tr._partir_audio(str(audio), tmp / "trozos")
            check("un audio largo se parte en varios trozos", len(trozos) == 3, str(len(trozos)))
            check("los trozos quedan en orden",
                  [x.name for x in trozos] == sorted(x.name for x in trozos))
            total = sum(duracion_audio_segundos(x) or 0 for x in trozos)
            check("no se pierde audio al cortar", abs(total - 50) < 2, f"{total:.1f}s de 50s")
        finally:
            tr.DURACION_TROZO_SEGUNDOS = original

        try:
            tr._verificar_parte("", str(audio), 900)
            check("una transcripcion vacia detiene la corrida", False)
        except ValueError as e:
            check("una transcripcion vacia detiene la corrida", True)
            check("el mensaje avisa que el audio no se perdio", "sigue donde estaba" in str(e))
            check("y sugiere cortarlo en partes", "partes" in str(e))

        try:
            tr._verificar_parte("palabra " * 300, str(audio), 900)
            check("una transcripcion normal pasa sin molestar", True)
        except ValueError:
            check("una transcripcion normal pasa sin molestar", False)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def probar_el_borrado_no_alcanza_tus_carpetas() -> None:
    """
    Cortar el audio en trozos implica borrarlos despues, y ese es el unico
    borrado que el sistema hace solo. Se prueba que NO PUEDA tocar otra cosa,
    no que "normalmente no la toque": perder grabaciones o notas de clase no
    tiene vuelta atras.
    """
    print("\n== el borrado de trozos no puede alcanzar nada mas ==")
    import tempfile as tf

    from orquestador import transcripcion as tr

    # 1. Solo borra carpetas creadas por el sistema, dentro del area temporal.
    propia = Path(tf.mkdtemp(prefix=tr.PREFIJO_TROZOS))
    (propia / "trozo.m4a").write_bytes(b"x")
    tr._borrar_carpeta_de_trozos(propia)
    check("borra su propia carpeta temporal", not propia.exists())

    # 2. Una carpeta temporal ajena (otro prefijo) no se toca.
    ajena = Path(tf.mkdtemp(prefix="algo_del_usuario_"))
    try:
        (ajena / "importante.md").write_text("no borrar", encoding="utf-8")
        tr._borrar_carpeta_de_trozos(ajena)
        check("no toca una carpeta temporal con otro nombre", ajena.exists())
    finally:
        shutil.rmtree(ajena, ignore_errors=True)

    # 3. Aunque el nombre calce, fuera del area temporal no se borra.
    fuera = Path(tf.mkdtemp()) / f"{tr.PREFIJO_TROZOS}falsa"
    fuera.mkdir(parents=True)
    simulada = fuera.parent / "Procesados"
    simulada.mkdir()
    (simulada / "clase.m4a").write_bytes(b"grabacion real")
    try:
        for objetivo in (Path.home() / "Documents", simulada, Path("/")):
            tr._borrar_carpeta_de_trozos(objetivo)
        check("no borra Documents, ni Procesados, ni la raiz",
              (Path.home() / "Documents").exists() and (simulada / "clase.m4a").exists())
    finally:
        shutil.rmtree(fuera.parent, ignore_errors=True)

    # 4. La ruta del audio original nunca entra en la carpeta de trozos.
    import inspect
    fuente = inspect.getsource(tr._transcribir_archivo)
    check("los trozos se crean en una carpeta temporal recien hecha",
          "tempfile.mkdtemp(prefix=PREFIJO_TROZOS)" in fuente)
    # Se mira el codigo, no la documentacion: el docstring justamente explica
    # que no se usa glob, y buscar la palabra ahi daba un falso positivo.
    cuerpo = inspect.getsource(tr._borrar_carpeta_de_trozos)
    cuerpo = cuerpo.split('"""')[-1]
    check("no se borra por patron ni comodin",
          "glob" not in cuerpo and "*" not in cuerpo, cuerpo.strip()[:60])
    check("se comprueba el area temporal antes de borrar",
          "gettempdir" in cuerpo and "is_relative_to" in cuerpo)


def probar_dos_clases_no_se_fusionan() -> None:
    """
    Dos clases reales terminaron en un mismo documento: "Desempeno
    organizacional 19.08.26" y "Econometria 20.08.26" se copiaron a Input la
    misma noche, las dos quedaron con mtime del 20, se agruparon como una sola
    clase de dos partes y mezclaron dos ramos distintos. El nombre tenia la
    fecha correcta y se estaba ignorando.
    """
    print("\n== dos clases distintas no se fusionan ==")
    import os
    from datetime import timedelta

    hoy = date.today()
    ayer = hoy - timedelta(days=1)

    tmp = Path(tempfile.mkdtemp())
    try:
        def crear(nombre: str, mtime_dia: date) -> Path:
            f = tmp / nombre
            f.write_bytes(b"audio")
            ts = __import__("time").mktime(mtime_dia.timetuple())
            os.utime(f, (ts, ts))
            return f

        # El caso real: dos clases de dias distintos, copiadas el mismo dia.
        a = crear(f"Desempeño organizacional {ayer:%d.%m.%y}.m4a", hoy)
        b = crear(f"Econometría {hoy:%d.%m.%y}.m4a", hoy)

        check("la fecha del nombre gana sobre la del archivo",
              deteccion._resolver_fecha_archivo(a) == ayer,
              str(deteccion._resolver_fecha_archivo(a)))
        grupos = deteccion.agrupar_por_fecha([a, b])
        check("quedan como dos clases separadas", len(grupos) == 2, str(list(grupos)))
        check("ninguna queda con dos audios",
              all(len(v) == 1 for v in grupos.values()))

        # Sin fecha en el nombre se sigue usando la del archivo.
        c = crear("grabacion sin fecha.m4a", ayer)
        check("sin fecha en el nombre manda el archivo",
              deteccion._resolver_fecha_archivo(c) == ayer)

        # Una fecha del futuro en el nombre no se cree.
        futuro = hoy + timedelta(days=30)
        d = crear(f"clase {futuro:%d.%m.%y}.m4a", hoy)
        check("una fecha futura en el nombre se descarta",
              deteccion._resolver_fecha_archivo(d) == hoy)

        # Un numero que parece fecha pero es de hace decadas tampoco.
        e = crear("Capitulo 1.2.99 repaso.m4a", ayer)
        check("un numero viejo que parece fecha se descarta",
              deteccion._resolver_fecha_archivo(e) == ayer)

        # Una clase partida de verdad por el estudiante si debe agruparse.
        f1 = crear(f"Marketing {ayer:%d.%m.%y} parte A.m4a", hoy)
        f2 = crear(f"Marketing {ayer:%d.%m.%y} parte B.m4a", hoy)
        g = deteccion.agrupar_por_fecha([f1, f2])
        check("dos archivos de la MISMA fecha si se agrupan juntos",
              len(g) == 1 and len(list(g.values())[0]) == 2)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def probar_las_notificaciones_no_se_pierden() -> None:
    """
    Una actualizacion de Homebrew dejo terminal-notifier sin poder entregar
    avisos, y el codigo no miraba el resultado: el estudiante habria dejado de
    recibir notificaciones sin enterarse. Ahora un fallo cae a la via nativa.
    """
    print("\n== las notificaciones no se pierden en silencio ==")
    from orquestador import notificaciones as n

    original_tn, original_as = n._con_terminal_notifier, n._con_applescript
    try:
        # Cuando la via preferida funciona, no se usa el respaldo.
        usados = []
        n._con_terminal_notifier = lambda *a, **k: usados.append("tn") or True
        n._con_applescript = lambda *a, **k: usados.append("as")
        n._notificar("t", "s", "m")
        check("si terminal-notifier entrega, no se duplica el aviso", usados == ["tn"], str(usados))

        # Cuando falla, el aviso igual llega.
        usados.clear()
        n._con_terminal_notifier = lambda *a, **k: usados.append("tn") or False
        n._notificar("t", "s", "m")
        check("si terminal-notifier falla, se usa la via nativa",
              usados == ["tn", "as"], str(usados))

        # Y si las dos fallan, no revienta el procesamiento.
        def explota(*a, **k):
            raise RuntimeError("sin notificaciones")
        n._con_terminal_notifier = lambda *a, **k: False
        n._con_applescript = explota
        try:
            n._notificar("t", "s", "m")
            check("un aviso que no se puede entregar no tumba la clase", True)
        except Exception as e:
            check("un aviso que no se puede entregar no tumba la clase", False,
                  f"propago {type(e).__name__}")
    finally:
        n._con_terminal_notifier, n._con_applescript = original_tn, original_as

    # Pase lo que pase, queda registrado en Estado.txt.
    import inspect
    check("todo aviso queda anotado antes de intentar entregarlo",
          "_anotar_estado" in inspect.getsource(n._notificar))
    check("el texto se escapa antes de pasarlo a AppleScript",
          "_escapar_applescript" in inspect.getsource(n._con_applescript))


def probar_error_de_sesion_se_explica() -> None:
    """
    La sesion de Claude Code caduco y el estudiante leyo "La skill termino con
    error: success", que no dice nada. La explicacion venia en el campo
    `result` del SDK y no se estaba mirando, asi que no habia forma de saber
    que lo que tocaba era volver a iniciar sesion.
    """
    print("\n== un error del modelo dice que paso de verdad ==")
    from orquestador.skill_runner import _es_sesion_caducada, describir_error_sdk

    class Falso:
        subtype = "success"
        is_error = True
        result = "Failed to authenticate: OAuth session expired and could not be refreshed"
        api_error_status = None
        terminal_reason = None
        errors = None

    texto = describir_error_sdk(Falso())
    check("el mensaje incluye la causa real", "OAuth session expired" in texto)
    check("y dice que hay que iniciar sesion de nuevo", "iniciar sesion" in texto.lower())
    check("y nombra el comando exacto", "node_modules/.bin/claude" in texto)
    check("y avisa que la transcripcion no se perdio", "transcripcion" in texto.lower())
    check("el subtype enganoso ya no va solo", texto.strip() != "subtype=success")

    check("reconoce una sesion caducada", _es_sesion_caducada(Falso.result))
    check("no confunde otro error con sesion caducada",
          not _es_sesion_caducada("API Error: 429 rate limit exceeded"))

    # Un error distinto se sigue describiendo, sin inventar el diagnostico.
    class Otro(Falso):
        result = "API Error: 529 overloaded"
    texto_otro = describir_error_sdk(Otro())
    check("otro error se muestra tal cual", "529" in texto_otro)
    check("y no sugiere iniciar sesion sin motivo",
          "node_modules/.bin/claude" not in texto_otro)


def probar_aviso_anki() -> None:
    print("\n== aviso cuando Anki esta cerrado ==")
    from orquestador import anki_connect, dialogo_anki

    original = anki_connect.verificar_conexion
    original_preg = dialogo_anki._preguntar
    try:
        anki_connect.verificar_conexion = lambda: True
        pregunto = []
        dialogo_anki._preguntar = lambda reintento=False: pregunto.append(1) or dialogo_anki.OPCION_SEGUIR
        sigue = dialogo_anki.confirmar_antes_de_empezar()
        check("con Anki abierto no molesta con ningun dialogo", sigue and not pregunto)

        anki_connect.verificar_conexion = lambda: False
        dialogo_anki._preguntar = lambda reintento=False: dialogo_anki.OPCION_SEGUIR
        check("'continuar sin Anki' deja procesar", dialogo_anki.confirmar_antes_de_empezar())

        # Dice que ya lo abrio pero sigue cerrado: se vuelve a comprobar en vez
        # de creerle, y despues de unos intentos se sigue igual.
        veces = []
        dialogo_anki._preguntar = lambda reintento=False: (veces.append(reintento)
                                                           or dialogo_anki.OPCION_YA_ABRI)
        check("'ya lo abri' con Anki aun cerrado no bloquea",
              dialogo_anki.confirmar_antes_de_empezar())
        check("reintenta y avisa que sigue sin detectarlo",
              len(veces) == dialogo_anki.INTENTOS_MAXIMOS and veces[1] is True, str(veces))
    finally:
        anki_connect.verificar_conexion = original
        dialogo_anki._preguntar = original_preg


def probar_titulo_nunca_falta() -> None:
    """
    El titulo era el unico campo que finalizar_clase leia sin respaldo, y una
    linea RESULTADO_ORQUESTADOR sin el costaba la clase entera con las notas ya
    escritas. Peor: el _skill.json quedaba guardado, asi que el reintento
    fallaba en el mismo punto para siempre.
    """
    print("\n== el titulo que falta no puede costar la clase ==")
    from orquestador.skill_runner import normalizar_resultado

    r = normalizar_resultado({"fuente": "a.md"}, "ECONOMETRIA")
    check("sin titulo, usa el ramo", r["titulo"] == "ECONOMETRIA")
    check("sin conceptos, deja la lista vacia", r["conceptos_repetidos"] == [])

    check("un titulo en blanco cuenta como ausente",
          normalizar_resultado({"titulo": "   "}, "R")["titulo"] == "R")
    check("un titulo que no es texto cuenta como ausente",
          normalizar_resultado({"titulo": 5}, "R")["titulo"] == "R")
    check("un titulo bueno no se toca",
          normalizar_resultado({"titulo": "Oferta y demanda"}, "R")["titulo"]
          == "Oferta y demanda")

    check("conceptos mal formados pasan a lista vacia",
          normalizar_resultado({"titulo": "T", "conceptos_repetidos": "no es lista"},
                               "R")["conceptos_repetidos"] == [])
    check("conceptos buenos no se tocan",
          normalizar_resultado({"titulo": "T", "conceptos_repetidos": ["a", "b"]},
                               "R")["conceptos_repetidos"] == ["a", "b"])

    # Un respaldo vacio dejaria el archivo terminado en " - ".
    sin_ramo = normalizar_resultado({}, "  ")["titulo"]
    check("sin ramo tampoco queda en blanco", sin_ramo.strip() != "")
    nombre = nombres.nombre_base(1, "2026-08-13", sin_ramo)
    check("el nombre de archivo no queda colgando de un guion",
          not nombre.rstrip().endswith("-"), nombre)


def probar_error_del_sdk_se_explica() -> None:
    """
    Un fallo HTTP de la API llega con subtype "success", asi que el mensaje que
    veia el estudiante era "La skill termino con error: success". Paso en vivo y
    costo reprocesar una clase entera.
    """
    print("\n== un error del SDK tiene que decir que paso ==")
    from orquestador.skill_runner import describir_error_sdk

    class _Falso:
        def __init__(self, **kw):
            self.subtype = "success"
            self.api_error_status = None
            self.terminal_reason = None
            self.errors = None
            self.__dict__.update(kw)

    texto = describir_error_sdk(_Falso(api_error_status=529))
    check("nombra el codigo HTTP cuando lo hay", "529" in texto, texto)
    check("y no se queda solo en el subtype engañoso", texto != "subtype=success", texto)

    texto = describir_error_sdk(_Falso(terminal_reason="max_turns"))
    check("dice por que termino la corrida", "max_turns" in texto, texto)

    texto = describir_error_sdk(_Falso(errors=["se cayo la conexion"]))
    check("incluye los errores que traiga", "se cayo la conexion" in texto, texto)

    # Un ResultMessage viejo puede no traer los campos nuevos del SDK.
    class _Viejo:
        subtype = "error_during_execution"

    check("no revienta si el SDK no trae los campos nuevos",
          "error_during_execution" in describir_error_sdk(_Viejo()))


def _falsear_query(mensajes):
    """Sustituto de claude_agent_sdk.query: emite mensajes ya preparados."""
    async def _query(*a, **kw):
        for m in mensajes:
            yield m
    return _query


def _mensajes_sdk(texto: str, is_error: bool, **extra):
    """Lo que emitiria el SDK en una corrida: la respuesta y el resultado."""
    from claude_agent_sdk import AssistantMessage, ResultMessage, TextBlock

    return [
        AssistantMessage(content=[TextBlock(text=texto)], model="falso"),
        ResultMessage(
            subtype="success",
            duration_ms=1,
            duration_api_ms=1,
            is_error=is_error,
            num_turns=3,
            session_id="sesion-1",
            usage={},
            total_cost_usd=0.0,
            **extra,
        ),
    ]


def probar_error_tardio_no_tira_el_trabajo() -> None:
    """
    Un 429 o un 529 al final de una corrida ya terminada no puede costar la
    clase. Paso en vivo dos veces: en una clase real (13 turnos perdidos y
    reprocesada entera) y en el ensayo, con las tres notas ya escritas.
    """
    print("\n== un error de la API al final no puede tirar el trabajo hecho ==")
    import asyncio

    from orquestador import skill_runner

    LINEA = ('RESULTADO_ORQUESTADOR: {"titulo": "Tema real", "fuente": "f.md", '
             '"aprendizaje": "a.md", "conceptos_repetidos": ["x"]}')

    avisos = []
    originales = (skill_runner.query, skill_runner.registrar_uso,
                  skill_runner.notificar_aviso)
    tmp = Path(tempfile.mkdtemp(prefix="pendientes_falsos_"))
    usar_dir_pendientes(tmp)
    skill_runner.registrar_uso = lambda *a, **kw: None
    skill_runner.notificar_aviso = lambda t, c: avisos.append((t, c))
    try:
        # 1. La skill alcanzo a reportar su trabajo y despues fallo la API.
        skill_runner.query = _falsear_query(
            _mensajes_sdk(LINEA, True, api_error_status=529)
        )
        r = asyncio.run(skill_runner.aplicar_skill(
            "t.txt", "ECONOMETRIA", "/tmp/vault", "slug1", "/tmp/vault/E"))
        check("la clase se salva pese al error de la API", r["titulo"] == "Tema real")
        check("el _skill.json queda guardado", (tmp / "slug1_skill.json").is_file())
        check("te avisa de que la API fallo", len(avisos) == 1, str(avisos))
        check("el aviso nombra el codigo HTTP", "529" in avisos[0][1], str(avisos))
        # Sin la terminal abierta, "2026-08-13_3e73c56e" no le dice nada a nadie.
        check("y nombra el ramo, no el slug interno",
              "ECONOMETRIA" in avisos[0][1] and "slug1" not in avisos[0][1], str(avisos))

        # 2. Sin error de la API no hay aviso que moleste.
        avisos.clear()
        skill_runner.query = _falsear_query(_mensajes_sdk(LINEA, False))
        asyncio.run(skill_runner.aplicar_skill(
            "t.txt", "ECONOMETRIA", "/tmp/vault", "slug2", "/tmp/vault/E"))
        check("una corrida limpia no dispara ningun aviso", avisos == [], str(avisos))

        # 3. Si ademas no hay nada aprovechable, revienta, pero explicando.
        skill_runner.query = _falsear_query(
            _mensajes_sdk("no reporte nada", True, api_error_status=429)
        )
        try:
            asyncio.run(skill_runner.aplicar_skill(
                "t.txt", "ECONOMETRIA", "/tmp/vault", "slug3", "/tmp/vault/E"))
            check("sin resultado utilizable si falla", False, "no revento")
        except ValueError as e:
            check("sin resultado utilizable si falla", True)
            check("y el error dice el codigo HTTP, no 'success'", "429" in str(e), str(e))

        # 4. La correccion tambien conserva su trabajo si la API falla al final.
        avisos.clear()
        corregida = ('RESULTADO_ORQUESTADOR: {"titulo": "Tema corregido", '
                     '"fuente": "f.md", "conceptos_repetidos": ["x"]}')
        skill_runner.query = _falsear_query(
            _mensajes_sdk(corregida, True, api_error_status=529)
        )
        previo = {"titulo": "Tema real", "session_id": "sesion-1"}
        r = asyncio.run(skill_runner.corregir_con_revision(
            previo, [{"que": "algo"}], "/tmp/vault", "slug4"))
        check("la correccion no se pierde por un error tardio",
              r["titulo"] == "Tema corregido", str(r))
        check("y queda marcada como corregida",
              r.get("corregido_tras_revision") is True, str(r))

        # 5. Si la correccion no reporta nada, se vuelve a lo anterior y avisa.
        avisos.clear()
        skill_runner.query = _falsear_query(
            _mensajes_sdk("no reporte nada", True, api_error_status=529)
        )
        r = asyncio.run(skill_runner.corregir_con_revision(
            previo, [{"que": "algo"}], "/tmp/vault", "slug5"))
        check("sin correccion confirmada se conserva lo anterior",
              r["titulo"] == "Tema real", str(r))
        check("y te avisa de que el documento va sin corregir",
              len(avisos) == 1 and "529" in avisos[0][1], str(avisos))
        check("ese aviso tambien nombra la clase, no el slug",
              "Tema real" in avisos[0][1] and "slug5" not in avisos[0][1], str(avisos))
    finally:
        (skill_runner.query, skill_runner.registrar_uso,
         skill_runner.notificar_aviso) = originales
        usar_dir_pendientes(None)
        shutil.rmtree(tmp, ignore_errors=True)


def probar_no_afirmar_que_el_profe_no_pidio_nada() -> None:
    """
    "El profesor no anuncio nada" y "no se pudo averiguar" no pueden verse
    igual. Antes se veian: un 429 devolvia dos listas vacias, el llamador las
    cacheaba en el _skill.json, y como solo se reintenta cuando el campo no
    existe, la clase quedaba afirmando en falso que no habia nada que estudiar.
    """
    print("\n== no afirmar que el profesor no pidio nada ==")
    import asyncio

    from orquestador import regenerar as rg

    avisos = []
    originales = (rg.query, rg.registrar_uso, rg.notificar_aviso)
    rg.registrar_uso = lambda *a, **kw: None
    rg.notificar_aviso = lambda t, c: avisos.append((t, c))
    try:
        LISTA = ('RESULTADO_LLAMADOS: {"avisos": [{"que": "prueba", "cuando": "", '
                 '"textual": "el jueves", "seguro": true}], "evaluacion": []}')
        VACIO = 'RESULTADO_LLAMADOS: {"avisos": [], "evaluacion": []}'

        # 1. Respuesta con anuncios de verdad.
        rg.query = _falsear_query(_mensajes_sdk(LISTA, False))
        r = asyncio.run(rg.extraer_llamados("t.txt", "ECONOMETRIA", "s1"))
        check("una respuesta con anuncios se devuelve tal cual",
              r and len(r["avisos"]) == 1, str(r))
        check("y no molesta con ningun aviso", avisos == [], str(avisos))

        # 2. El profesor de verdad no anuncio nada. Es respuesta legitima.
        rg.query = _falsear_query(_mensajes_sdk(VACIO, False))
        r = asyncio.run(rg.extraer_llamados("t.txt", "ECONOMETRIA", "s2"))
        check("dos listas vacias comprobadas son una respuesta valida",
              r == {"avisos": [], "evaluacion": []}, str(r))
        check("tampoco avisa nada en ese caso", avisos == [], str(avisos))

        # 3. La API fallo y ademas no hubo linea: aqui si es "no se sabe".
        rg.query = _falsear_query(
            _mensajes_sdk("no alcance a responder", True, api_error_status=429)
        )
        r = asyncio.run(rg.extraer_llamados("t.txt", "ECONOMETRIA", "s3"))
        check("un fallo de la API no se hace pasar por lista vacia", r is None, str(r))
        check("y te avisa nombrando el codigo HTTP",
              len(avisos) == 1 and "429" in avisos[0][1], str(avisos))
        check("el aviso aclara que no es lo mismo que no haber anuncios",
              "no significa" in avisos[0][1], str(avisos))

        # 4. El modelo contesto pero sin la linea, sin error de la API.
        avisos.clear()
        rg.query = _falsear_query(_mensajes_sdk("no reporte nada", False))
        r = asyncio.run(rg.extraer_llamados("t.txt", "ECONOMETRIA", "s4"))
        check("una respuesta sin la linea tampoco se hace pasar por vacia", r is None)
        check("y tambien avisa", len(avisos) == 1, str(avisos))

        # 5. La API fallo DESPUES de que el modelo ya habia reportado. La
        # respuesta vale: emitir la linea es la prueba de que el trabajo se
        # hizo. Mismo criterio que en aplicar_skill y corregir_con_revision, y
        # esta prueba existe para que nadie lo "corrija" en sentido contrario.
        avisos.clear()
        rg.query = _falsear_query(_mensajes_sdk(LISTA, True, api_error_status=529))
        r = asyncio.run(rg.extraer_llamados("t.txt", "ECONOMETRIA", "s5"))
        check("un error posterior al reporte no descarta la respuesta",
              r and len(r["avisos"]) == 1, str(r))
    finally:
        (rg.query, rg.registrar_uso, rg.notificar_aviso) = originales


def probar_lo_no_comprobado_no_se_cachea() -> None:
    """
    El dano real no era devolver vacio, era guardarlo: la condicion para
    reintentar es que el campo no exista, asi que un vacio cacheado congela la
    mentira para siempre.
    """
    print("\n== lo que no se pudo comprobar no se guarda ==")
    import asyncio
    import json as _json

    from orquestador import regenerar as rg

    tmp = Path(tempfile.mkdtemp(prefix="regenerar_falso_"))
    usar_dir_pendientes(tmp)
    originales = (rg.extraer_llamados, rg.generar_docx, rg._leer_nota)
    rg.generar_docx = lambda *a, **kw: tmp / "falso.docx"
    rg._leer_nota = lambda ruta, vault: "texto"
    try:
        cfg = {"rutas": {"vault_obsidian": str(tmp), "output": str(tmp)}}

        def preparar(slug):
            (tmp / f"{slug}.json").write_text(_json.dumps({
                "ramo": "ECONOMETRIA", "fecha": "2026-08-13", "numero_clase": 2,
                "archivo_texto": str(tmp / "t.txt"), "archivos_originales": [],
            }), encoding="utf-8")
            (tmp / f"{slug}_skill.json").write_text(_json.dumps({
                "titulo": "Tema", "fuente": "f.md", "conceptos_repetidos": [],
            }), encoding="utf-8")

        # No se pudo averiguar: el _skill.json no debe quedar con "llamados".
        preparar("sinsaber")
        rg.extraer_llamados = lambda *a, **kw: _corutina(None)
        asyncio.run(rg.regenerar("sinsaber", cfg))
        guardado = _json.loads((tmp / "sinsaber_skill.json").read_text(encoding="utf-8"))
        check("un resultado no comprobado no se guarda",
              "llamados" not in guardado, str(guardado.get("llamados")))
        check("asi la proxima regeneracion lo vuelve a intentar",
              guardado.get("llamados") is None)

        # Comprobado: si se guarda, para no volver a pagar la lectura.
        preparar("sabido")
        rg.extraer_llamados = lambda *a, **kw: _corutina({"avisos": [], "evaluacion": []})
        asyncio.run(rg.regenerar("sabido", cfg))
        guardado = _json.loads((tmp / "sabido_skill.json").read_text(encoding="utf-8"))
        check("un vacio comprobado si se guarda",
              guardado.get("llamados") == {"avisos": [], "evaluacion": []}, str(guardado))
    finally:
        (rg.extraer_llamados, rg.generar_docx, rg._leer_nota) = originales
        usar_dir_pendientes(None)
        shutil.rmtree(tmp, ignore_errors=True)


def probar_revision_fallida_no_dice_aprobado() -> None:
    """
    El revisor es la unica barrera antes de que el material entre a Anki. Una
    revision que no llego a correr quedaba guardada como "aprobado", o sea con
    cara de comprobada y limpia, y el campo que guardaba la verdad no lo leia
    nadie.
    """
    print("\n== una revision que no corrio no puede decir 'aprobado' ==")
    import asyncio

    from orquestador import revisor as rv

    originales = (rv.query, rv.registrar_uso)
    tmp = Path(tempfile.mkdtemp(prefix="revision_falsa_"))
    usar_dir_pendientes(tmp)
    rv.registrar_uso = lambda *a, **kw: None
    try:
        previo = {"fuente": "f.md", "aprendizaje": "a.md"}

        # El revisor no alcanzo a emitir su linea.
        rv.query = _falsear_query(_mensajes_sdk("no alcance", True, api_error_status=529))
        r = asyncio.run(rv.revisar("t.txt", previo, "ECONOMETRIA", "/tmp/v", "s1"))
        check("el veredicto no es 'aprobado'", r["veredicto"] != "aprobado", str(r))
        check("dice explicitamente que no se reviso",
              r["veredicto"] == "no_revisado", str(r))
        check("y guarda el motivo con el codigo HTTP",
              "529" in r["revision_fallida"], str(r))
        # finalizar_clase decide corregir comparando contra "corregir": un
        # veredicto nuevo no puede disparar una correccion por accidente.
        check("no dispara correccion", r["veredicto"] != "corregir")
        check("y no aporta hallazgos falsos", r["hallazgos"] == [])

        # Una revision que si corrio sigue funcionando igual que siempre.
        buena = ('RESULTADO_REVISION: {"veredicto": "aprobado", "hallazgos": []}')
        rv.query = _falsear_query(_mensajes_sdk(buena, False))
        r = asyncio.run(rv.revisar("t.txt", previo, "ECONOMETRIA", "/tmp/v", "s2"))
        check("una revision real si puede aprobar", r["veredicto"] == "aprobado", str(r))
        check("y no queda marcada como fallida",
              not r.get("revision_fallida"), str(r))
    finally:
        (rv.query, rv.registrar_uso) = originales
        usar_dir_pendientes(None)
        shutil.rmtree(tmp, ignore_errors=True)


def probar_una_sola_comprobacion_de_vault() -> None:
    """
    La comprobacion que impide que una ruta colada en la transcripcion meta
    cualquier archivo del disco en el .docx estaba duplicada en dos modulos.
    """
    print("\n== la frontera del vault se comprueba en un solo lugar ==")
    from orquestador import finalizar_clase as fc
    from orquestador import regenerar as rg

    vault = Path(tempfile.mkdtemp(prefix="notas_falsas_"))
    try:
        (vault / "buena.md").write_text("contenido de la nota", encoding="utf-8")
        fuera = Path(tempfile.mkdtemp(prefix="fuera_")) / "secreto.md"
        fuera.parent.mkdir(parents=True, exist_ok=True)
        fuera.write_text("esto no puede terminar en el docx", encoding="utf-8")

        for nombre, leer in (("finalizar_clase", fc._leer_nota), ("regenerar", rg._leer_nota)):
            check(f"{nombre}: lee una nota de dentro del vault",
                  leer(str(vault / "buena.md"), str(vault)) == "contenido de la nota")
            check(f"{nombre}: no lee nada de fuera del vault",
                  leer(str(fuera), str(vault)) == "")
            check(f"{nombre}: una ruta vacia no revienta", leer(None, str(vault)) == "")
            check(f"{nombre}: un archivo que no existe da vacio",
                  leer(str(vault / "no_existe.md"), str(vault)) == "")

        shutil.rmtree(fuera.parent, ignore_errors=True)
    finally:
        shutil.rmtree(vault, ignore_errors=True)


async def _corutina(valor):
    """Envuelve un valor para sustituir una funcion async en las pruebas."""
    return valor


def _gate_deja_pasar(gate, ruta, clave: str = "file_path") -> bool:
    """Corre el hook y dice si autorizo la ruta. Devolver {} es autorizar."""
    import asyncio

    respuesta = asyncio.run(gate({"tool_input": {clave: str(ruta)}}, "id-1", None))
    return respuesta == {}


def _gate_motivo(gate, ruta) -> str:
    import asyncio

    respuesta = asyncio.run(gate({"tool_input": {"file_path": str(ruta)}}, "id-1", None))
    return respuesta["hookSpecificOutput"]["permissionDecisionReason"]


def probar_gate_de_rutas() -> None:
    """
    El gate es el unico freno real de la corrida automatizada: corre con
    permission_mode "bypassPermissions", asi que sin el Write podria escribir en
    cualquier parte del disco. Nunca habia tenido pruebas.
    """
    print("\n== el gate de rutas de la corrida automatizada ==")
    from orquestador.skill_runner import construir_gate_de_rutas

    # El prefijo evita a proposito la palabra "vault": el mensaje de denegacion
    # repite la ruta denegada, y con ella dentro no se podria comprobar que el
    # mensaje no nombre un vault entre las carpetas que si autorizo.
    vault = Path(tempfile.mkdtemp(prefix="notas_falsas_"))
    try:
        gate = construir_gate_de_rutas(RAIZ, vault)

        check("deja leer la transcripcion, que vive en el proyecto",
              _gate_deja_pasar(gate, RAIZ / "orquestador" / "transcripciones_pendientes" / "x.txt"))
        check("deja escribir la nota, que va en el vault",
              _gate_deja_pasar(gate, vault / "ECONOMETRIA" / "Clase.md"))
        check("deniega cualquier otra parte del disco",
              not _gate_deja_pasar(gate, "/etc/passwd"))
        check("deniega una carpeta vecina del proyecto",
              not _gate_deja_pasar(gate, RAIZ.parent / "Otra cosa" / "a.md"))
        check("deniega el home entero",
              not _gate_deja_pasar(gate, "~/nota.md"))

        # Glob y Grep no usan "file_path" sino "path": si el gate mirara solo
        # una de las dos claves, la otra pasaria sin control.
        check("tambien gatea la clave 'path' de Glob y Grep",
              not _gate_deja_pasar(gate, "/etc", clave="path"))

        # Una ruta con .. se resuelve antes de comparar, si no el gate se
        # esquivaria escribiendo hacia arriba desde una carpeta autorizada.
        check("un .. no escapa de una raiz autorizada",
              not _gate_deja_pasar(gate, RAIZ / ".." / ".." / "etc" / "passwd"))

        # Sin ruta no hay nada que validar (ej. una herramienta sin archivo).
        import asyncio
        check("una llamada sin ruta no se bloquea",
              asyncio.run(gate({"tool_input": {}}, "id-1", None)) == {})

        # La etapa que solo extrae los llamados a la accion corre con una sola
        # raiz: el vault no tiene por que estar a su alcance (ver regenerar.py).
        solo_proyecto = construir_gate_de_rutas(RAIZ)
        check("con una sola raiz, el vault queda fuera",
              not _gate_deja_pasar(solo_proyecto, vault / "nota.md"))
        motivo = _gate_motivo(solo_proyecto, vault / "nota.md")
        check("y el mensaje no inventa un vault que no autorizo",
              "vault" not in motivo.lower(), motivo)
        check("el mensaje dice cual es la raiz permitida", str(RAIZ) in motivo, motivo)

        # Los intermedios del ensayo viven en el temp del sistema, fuera de las
        # raices. Hoy no los lee el modelo (la metadata trae la ruta real del
        # proyecto, ver finalizar_clase), pero si algun dia se movieran ahi,
        # esta prueba lo dice en vez de fallar en vivo a mitad de una clase.
        sandbox = Path(tempfile.mkdtemp(prefix="ensayo_falso_"))
        try:
            gate_ensayo = construir_gate_de_rutas(RAIZ, sandbox / "vault")
            check("un intermedio del sandbox de ensayo no esta autorizado",
                  not _gate_deja_pasar(gate_ensayo, sandbox / "pendientes" / "x.txt"))
        finally:
            shutil.rmtree(sandbox, ignore_errors=True)

        # Un gate sin raices denegaria todo y la corrida fallaria en el primer
        # Read, lejos de la llamada que se equivoco.
        try:
            construir_gate_de_rutas()
            check("un gate sin raices se rechaza al construirlo", False, "no reventó")
        except ValueError:
            check("un gate sin raices se rechaza al construirlo", True)
    finally:
        shutil.rmtree(vault, ignore_errors=True)


if __name__ == "__main__":
    probar_nombres()
    probar_deteccion()
    probar_flashcards()
    probar_docx()
    probar_aislamiento_del_ensayo()
    probar_archivado_no_destructivo()
    probar_dialogo_nunca_descarta_solo()
    probar_el_nombre_del_archivo_manda_sobre_el_dia()
    probar_bitacora_deshace_todo()
    probar_bitacora_no_borra_lo_ajeno()
    probar_seccion_critica()
    probar_formulas()
    probar_contexto_va_primero()
    probar_formulas_en_tablas()
    probar_matematica_dentro_de_frases()
    probar_mapa_y_secciones()
    probar_llamados_a_la_accion()
    probar_documento_no_se_repite()
    probar_canales_de_enfasis()
    probar_parrafos_y_justificado()
    probar_audio_largo_se_corta_solo()
    probar_el_borrado_no_alcanza_tus_carpetas()
    probar_dos_clases_no_se_fusionan()
    probar_las_notificaciones_no_se_pierden()
    probar_error_de_sesion_se_explica()
    probar_aviso_anki()
    probar_titulo_nunca_falta()
    probar_gate_de_rutas()
    probar_error_del_sdk_se_explica()
    probar_error_tardio_no_tira_el_trabajo()
    probar_no_afirmar_que_el_profe_no_pidio_nada()
    probar_lo_no_comprobado_no_se_cachea()
    probar_revision_fallida_no_dice_aprobado()
    probar_una_sola_comprobacion_de_vault()

    print()
    if fallos:
        print(f"FALLARON {len(fallos)}: " + ", ".join(fallos))
        raise SystemExit(1)
    print("Todas las pruebas pasaron.")
