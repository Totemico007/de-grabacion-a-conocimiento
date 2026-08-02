"""
Renderiza formulas matematicas para el .docx.

Dos caminos, segun donde va la formula:

- **Formula destacada** (sola en su linea, entre `$$`): se dibuja como imagen
  con matplotlib y se inserta centrada. Queda con tipografia matematica de
  verdad: barra de fraccion, radical que se estira sobre lo que cubre,
  sombrero sobre la p, subindices bien puestos.
- **Formula corta dentro de una frase** (entre `$`): se escribe como texto con
  subindices y superindices reales. Una imagen ahi quedaria desalineada con el
  renglon y con un tamano que no acompana al texto.

Por que no ecuaciones nativas de Word. Word usa OMML, un XML propio que
python-docx no sabe generar: habria que armarlo a mano por cada formula, es
fragil, y una mal formada deja el .docx sin abrir. La imagen se ve igual de
bien y no puede romper el documento.

Por que no cuesta tokens. El modelo escribe la formula en LaTeX, que es lo que
ya escribiria de todos modos. El dibujo lo hace matplotlib aca, en el Mac.

Si matplotlib no esta o falla, se cae solo al modo texto. Una formula fea es
mucho mejor que una clase sin documento.
"""
import re
import tempfile
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

MARCA_DESTACADA = "$$"

# Un token es: _{...}, _(...), ^{...}, ^(...), _c, ^c, o texto normal.
# Se aceptan tambien los parentesis porque es lo que el modelo escribe cuando
# no esta usando LaTeX (visto en vivo: `z_(1-α/2)`), y era mejor entenderlo que
# mostrarlo crudo.
_PATRON = re.compile(r"(_\{[^}]*\}|\^\{[^}]*\}|_\([^)]*\)|\^\([^)]*\)|_[^\s_^{}()]|\^[^\s_^{}()])")

# Formula corta metida dentro de una frase: `la media $x̄$ se calcula...`.
# Se delimita a proposito: sin delimitador, cualquier guion bajo suelto en una
# palabra normal se leeria como subindice y destrozaria el parrafo.
_INLINE = re.compile(r"\$([^$]+)\$")

# Tamano de letra al dibujar, y resolucion. Mas dpi es mas nitido y pesa poco:
# una formula tipica ronda los 7 KB.
TAMANO_FUENTE = 20
DPI = 220
ANCHO_MAXIMO_PULGADAS = 6.0
# Dentro de una celda de tabla hay mucho menos espacio: una imagen a tamano
# natural desbordaria la columna y descuadraria la tabla entera.
ANCHO_MAXIMO_CELDA_PULGADAS = 2.3

_carpeta_temporal: Path | None = None


def _carpeta() -> Path:
    global _carpeta_temporal
    if _carpeta_temporal is None:
        _carpeta_temporal = Path(tempfile.mkdtemp(prefix="formulas_"))
    return _carpeta_temporal


def limpiar_temporales() -> None:
    """Las imagenes quedan incrustadas dentro del .docx, asi que los archivos
    sueltos ya no hacen falta una vez guardado."""
    global _carpeta_temporal
    if _carpeta_temporal is not None:
        import shutil
        shutil.rmtree(_carpeta_temporal, ignore_errors=True)
        _carpeta_temporal = None


def es_formula_destacada(linea: str) -> bool:
    limpia = linea.strip()
    return limpia.startswith(MARCA_DESTACADA) and limpia.endswith(MARCA_DESTACADA) and len(limpia) > 4


def texto_de_formula_destacada(linea: str) -> str:
    return linea.strip()[len(MARCA_DESTACADA):-len(MARCA_DESTACADA)].strip()


def partir_por_formulas_inline(texto: str) -> list[tuple[str, bool]]:
    """Parte el texto en tramos (contenido, es_formula), respetando el orden."""
    partes = []
    ultimo = 0
    for m in _INLINE.finditer(texto):
        if m.start() > ultimo:
            partes.append((texto[ultimo:m.start()], False))
        partes.append((m.group(1), True))
        ultimo = m.end()
    if ultimo < len(texto):
        partes.append((texto[ultimo:], False))
    return partes or [(texto, False)]


def escribir_en_parrafo(paragraph, texto: str, cursiva: bool = True) -> None:
    """Modo texto: subindices y superindices reales, sin imagen."""
    for parte in _PATRON.split(texto):
        if not parte:
            continue
        if parte.startswith("_") or parte.startswith("^"):
            contenido = parte[1:]
            if contenido[:1] in ("{", "("):
                contenido = contenido[1:-1]
            run = paragraph.add_run(contenido)
            run.font.subscript = parte.startswith("_")
            run.font.superscript = parte.startswith("^")
        else:
            run = paragraph.add_run(parte)
        run.italic = cursiva


def _renderizar_imagen(latex: str, ancho_maximo: float = ANCHO_MAXIMO_PULGADAS) -> tuple[Path, float] | None:
    """Dibuja la formula y devuelve (ruta, ancho en pulgadas). None si no se pudo."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        fig = plt.figure(figsize=(0.01, 0.01))
        fig.text(0, 0, f"${a_latex(latex)}$", fontsize=TAMANO_FUENTE)
        destino = _carpeta() / f"f{abs(hash(latex))}.png"
        fig.savefig(destino, dpi=DPI, bbox_inches="tight", pad_inches=0.12, transparent=True)
        plt.close(fig)

        from PIL import Image
        with Image.open(destino) as img:
            ancho = min(img.width / DPI, ancho_maximo)
        return destino, ancho
    except Exception:
        return None


def agregar_formula_destacada(doc, latex: str) -> None:
    """Formula en su propio parrafo, centrada. Intenta imagen y si no, texto."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    imagen = _renderizar_imagen(latex)
    if imagen is not None:
        ruta, ancho = imagen
        p.add_run().add_picture(str(ruta), width=Inches(ancho))
        return

    escribir_en_parrafo(p, latex)
    for run in p.runs:
        run.font.size = Pt(13)


def agregar_formula_en_celda(celda, latex: str) -> bool:
    """
    Dibuja la formula dentro de una celda de tabla.

    Existe porque las tablas se armaban con `celda.text = ...`, o sea texto
    plano sin ningun formato: justo donde mas formulas hay (la tabla de "que
    formula uso en cada caso") era donde peor se veian. Devuelve False si no se
    pudo dibujar, para que quien llame escriba el texto con subindices.
    """
    imagen = _renderizar_imagen(latex, ANCHO_MAXIMO_CELDA_PULGADAS)
    if imagen is None:
        return False
    ruta, ancho = imagen
    p = celda.paragraphs[0]
    p.add_run().add_picture(str(ruta), width=Inches(ancho))
    return True


def parece_solo_formula(texto: str) -> bool:
    """
    Una celda que es enteramente una formula se dibuja; una que es una frase
    con un simbolo suelto, no. Sin esta distincion, una celda de texto normal
    terminaria convertida en imagen y dejaria de poder buscarse o copiarse.
    """
    limpio = texto.strip()
    if not limpio or len(limpio) > 90:
        return False
    if limpio.startswith("$") and limpio.endswith("$"):
        return True
    # Sin delimitadores: se considera formula si casi no tiene palabras y si
    # trae senales matematicas claras.
    senales = any(s in limpio for s in ("±", "√", "σ", "μ", "α", "Σ", "·", "^", "_", "≤", "≥", "≈"))
    palabras_largas = sum(1 for palabra in re.findall(r"[A-Za-zÁÉÍÓÚáéíóúÑñ]{4,}", limpio))
    return senales and palabras_largas <= 1


# Traduccion de la notacion Unicode a LaTeX, para poder DIBUJAR una formula que
# el modelo no escribio en LaTeX.
#
# Hace falta porque el modelo no siempre usa LaTeX: dentro de una tabla escribe
# `x̄ ± z_(1-α/2)·σ/√n`, que es perfectamente legible para el, pero matplotlib
# no sabe dibujarlo. Antes esas formulas se mostraban crudas, con el guion bajo
# y el parentesis a la vista, justo en la tabla donde mas formulas hay.
_GRIEGAS = {
    "α": r"\alpha", "β": r"\beta", "γ": r"\gamma", "δ": r"\delta",
    "ε": r"\epsilon", "θ": r"\theta", "λ": r"\lambda", "μ": r"\mu",
    "π": r"\pi", "ρ": r"\rho", "σ": r"\sigma", "τ": r"\tau",
    "φ": r"\phi", "χ": r"\chi", "ω": r"\omega",
    "Δ": r"\Delta", "Σ": r"\sum", "Ω": r"\Omega", "Φ": r"\Phi",
}
_SIMBOLOS = {
    "±": r"\pm", "·": r"\cdot", "×": r"\times", "÷": r"\div",
    "≈": r"\approx", "≤": r"\leq", "≥": r"\geq", "≠": r"\neq",
    "∞": r"\infty", "∈": r"\in", "→": r"\to", "∫": r"\int",
}
# Letra + acento combinante: x̄ es "x" seguido de U+0304.
_ACENTOS = {"\u0304": "bar", "\u0302": "hat", "\u0303": "tilde", "\u0307": "dot"}


def a_latex(texto: str) -> str:
    """Convierte notacion Unicode a LaTeX. Si ya viene en LaTeX, lo deja igual."""
    if "\\" in texto:
        return texto

    # Acentos combinantes primero: dependen del caracter que los precede.
    resultado = []
    for caracter in texto:
        if caracter in _ACENTOS and resultado:
            resultado[-1] = f"\\{_ACENTOS[caracter]}{{{resultado[-1]}}}"
        else:
            resultado.append(caracter)
    texto = "".join(resultado)

    for unicode_, latex in {**_GRIEGAS, **_SIMBOLOS}.items():
        texto = texto.replace(unicode_, latex + " ")

    # Raices: √(algo) y √algo
    texto = re.sub(r"√\(([^)]*)\)", r"\\sqrt{\1}", texto)
    texto = re.sub(r"√\s*([A-Za-z0-9]+)", r"\\sqrt{\1}", texto)

    # Subindices y superindices con parentesis a la forma con llaves.
    texto = re.sub(r"_\(([^)]*)\)", r"_{\1}", texto)
    texto = re.sub(r"\^\(([^)]*)\)", r"^{\1}", texto)

    # Las palabras sueltas se escriben derechas, no en cursiva matematica.
    texto = re.sub(r"(?<![\\{])\b([A-Za-zÁÉÍÓÚáéíóúÑñ]{3,})\b", r"\\mathrm{\1}", texto)
    return texto
