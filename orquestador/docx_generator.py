"""
Etapa 6 (parte 2): arma el .docx de respaldo de una clase, con diseno pensado
para aprender (tipografia clara, encabezados con color, puntos clave
destacados), no un volcado de texto plano. Combina:
- la nota de fuente limpia y la nota de aprendizaje que genero la skill
  (etapa 4), leidas desde las rutas que la skill reporto
- los 5 conceptos mas repetidos por el profesor (etapa 5)

Se guarda en Output/[Ramo]/<mismo nombre de clase que el audio archivado>.docx
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import OxmlElement

from . import formulas, mapa_visual
from .nombres import nombre_base

# Un solo color de acento para todos los encabezados. Antes habia un tono por
# nivel, pero el nivel ya lo dice el tamano de la letra: el segundo color solo
# gastaba presupuesto de atencion (ver references/diseno-documento.md).
COLOR_ACENTO = RGBColor(0x1F, 0x4E, 0x5F)
COLOR_TABLA_HEADER = "2E7486"

# Los dos unicos canales de enfasis del documento, cada uno con un significado
# fijo. Destacar funciona porque es escaso: el metaanalisis de Schneider (2018,
# 103 estudios) da g = 0,53 en retencion, y ese efecto se pierde si todo
# resalta. Por eso no hay un tercer color esperando que alguien lo use.
COLOR_EXAMEN_FONDO = "FFF2CC"      # el profesor dijo que esto entra
COLOR_VERIFICAR_FONDO = "F2F2F2"   # esto puede estar mal
COLOR_CONTEXTO_FONDO = "EAF3F7"    # esto no sale de la clase

# Margenes. Con cuerpo 11, los 2,5 cm que trae Word dejan la linea en unos 95
# caracteres y el ojo pierde el renglon al volver a la izquierda. 3,5 cm la
# dejan cerca de 75, dentro del rango de 50 a 75 que recomienda la
# investigacion de legibilidad.
MARGEN_LATERAL_CM = 3.5
INTERLINEADO = 1.35

# Texto justificado en el cuerpo. Se puede apagar cambiando esto a False, sin
# tocar nada mas.
#
# Los cortes raros que se vieron al principio no venian de aqui: venian de que
# cada renglon de la nota se convertia en un parrafo suelto (ver
# _unir_lineas_de_parrafo). Justificar solo los hacia mas visibles.
#
# Lo que si es real: las guias de accesibilidad (WCAG 1.4.8) desaconsejan
# justificar, porque el espaciado desigual entre palabras forma "rios" blancos
# que cuestan mas de seguir, sobre todo con dislexia. Eso se mitiga con la
# particion de palabras que se activa mas abajo y con la linea a 75 caracteres,
# pero no desaparece del todo. Queda como preferencia, no como recomendacion.
JUSTIFICADO = True

# Partes del kit de repaso pensadas para repaso espaciado en varios dias, que
# no aportan cuando el estudiante tiene solo unas horas antes de la prueba (pidio
# sacarlas del .docx). La nota de Obsidian las sigue teniendo completas, por
# si las usa mas adelante para repaso de largo plazo.
SECCIONES_A_OMITIR_EN_DOCX = [
    "preguntas de repaso",
    "plan de repaso",
    # La sesion por bloques de tiempo dice COMO estudiar; lo que el estudiante
    # necesita en el .docx es la materia ya digerida y con ejemplos, que va en
    # "La materia" (ver SKILL.md). La sesion sigue completa en la nota de
    # Obsidian para quien quiera organizarse con ella.
    "sesión de estudio",
    "sesion de estudio",
    # Los llamados a la accion se dibujan aparte y de primeros, desde el JSON
    # estructurado que reporta la skill, para poder destacarlos y mostrar la
    # cita textual. La nota tambien los trae como seccion, para Obsidian, y esa
    # copia se salta aca para no decir lo mismo dos veces.
    "lo que el profesor pidió",
    "lo que el profesor pidio",
    # Los conceptos repetidos se dibujan como tabla, tambien desde el JSON. La
    # nota los repite como seccion para Obsidian, y sin esto salian dos veces:
    # una como tabla y otra como lista, a tres paginas de distancia.
    "conceptos que el profesor más repitió",
    "conceptos que el profesor mas repitio",
    "conceptos mas repetidos",
    "conceptos más repetidos",
]

# Compatibilidad con notas escritas antes de la fusion de secciones. En esas
# notas los pasos 1, 2 y 5 del metodo estan separados y explican los mismos
# conceptos tres veces. Las notas nuevas ya vienen fusionadas por la skill (ver
# SKILL.md), asi que esto solo se aplica cuando se detecta la version vieja: la
# senal es que exista la seccion desarrollada Y ademas las otras dos.
SECCIONES_FUSIONADAS = ["conceptos centrales", "para enseñarlo desde cero", "para enseñarlo"]
SECCION_DESARROLLADA = ["materia lista para estudiar", "la materia"]


def _secciones_a_omitir(texto_aprendizaje: str) -> list[str]:
    texto = texto_aprendizaje.lower()
    tiene_desarrollada = any(s in texto for s in SECCION_DESARROLLADA)
    tiene_repetidas = any(s in texto for s in SECCIONES_FUSIONADAS)
    if tiene_desarrollada and tiene_repetidas:
        return SECCIONES_A_OMITIR_EN_DOCX + SECCIONES_FUSIONADAS
    return SECCIONES_A_OMITIR_EN_DOCX


# Secciones de la nota de fuente que sí entran al .docx. El desarrollo
# cronologico de la clase no esta: eso es la misma materia contada otra vez, en
# otro orden, y ahora vive desarrollada dentro de "La materia" (los casos y
# ejemplos del profesor incluidos). Lo que queda aca es lo que la materia no
# cubre: las definiciones textuales, lo reconstruido y lo que falta.
SECCIONES_FUENTE_EN_DOCX = [
    "definiciones y ejemplos",
    "gráficos",
    "graficos",
    "huecos",
]


def _sombrear_parrafo(paragraph, color_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    paragraph._p.get_or_add_pPr().append(shd)


def _barra_izquierda(paragraph, color_hex: str) -> None:
    """
    Barra vertical al costado del parrafo. Es lo que hace que un bloque
    destacado se vea como bloque incluso al hojear rapido: el sombreado solo se
    confunde con una tabla, la barra dice "esto es aparte".
    """
    pPr = paragraph._p.get_or_add_pPr()
    bordes = OxmlElement("w:pBdr")
    izq = OxmlElement("w:left")
    izq.set(qn("w:val"), "single")
    izq.set(qn("w:sz"), "18")
    izq.set(qn("w:space"), "8")
    izq.set(qn("w:color"), color_hex)
    bordes.append(izq)
    pPr.append(bordes)


def _bloque_destacado(doc: Document, etiqueta: str, texto: str, color_fondo: str,
                      color_barra: str) -> None:
    """Un parrafo con fondo, barra lateral y una etiqueta corta en negrita."""
    p = doc.add_paragraph()
    _sombrear_parrafo(p, color_fondo)
    _barra_izquierda(p, color_barra)
    if etiqueta:
        run = p.add_run(f"{etiqueta}  ")
        run.bold = True
    _agregar_texto_con_negritas(p, texto)
    return p


def _sombrear_celda(celda, color_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color_hex)
    celda._tc.get_or_add_tcPr().append(shd)


def _agregar_texto_con_negritas(paragraph, texto: str) -> None:
    partes = re.split(r"(\*\*.+?\*\*)", texto)
    for parte in partes:
        if not parte:
            continue
        if parte.startswith("**") and parte.endswith("**"):
            formulas.escribir_prosa_con_matematica(paragraph, parte[2:-2], negrita=True)
        else:
            # El texto puede traer formulas cortas delimitadas ($x̄$) y tambien
            # subindices sueltos en medio de la frase, sin delimitar, que es lo
            # que el modelo escribe en la practica (ver formulas.py).
            for trozo, es_formula in formulas.partir_por_formulas_inline(parte):
                if es_formula:
                    formulas.escribir_en_parrafo(paragraph, trozo)
                elif trozo:
                    formulas.escribir_prosa_con_matematica(paragraph, trozo)


# El orden en que se estudia, que no es el orden en que la nota fue escrita:
# primero se entiende la materia y despues uno se pone a prueba. Las notas
# traen las preguntas antes porque asi van los pasos del metodo MIT, pero leer
# las preguntas antes de la materia no ayuda a nadie.
#
# Las secciones que no esten en esta lista se mantienen en su orden, al final.
ORDEN_SECCIONES = [
    "materia",
    "mapa visual",
    "preguntas",
    "respuestas modelo",
    "kit de repaso",
]


def _quitar_numero(titulo: str) -> str:
    """
    Quita el "3. " de "3. Diez preguntas".

    La numeracion viene de los pasos del metodo MIT y deja de tener sentido en
    cuanto se fusiona o se salta una seccion: el documento quedaba empezando en
    el punto 3. El orden ya lo dice el orden.
    """
    return re.sub(r"^\s*\d+\.\s+", "", titulo)


def _reordenar_secciones(md_texto: str) -> str:
    """Reordena los bloques de nivel 2 segun ORDEN_SECCIONES."""
    lineas = md_texto.splitlines()
    cabecera: list[str] = []
    bloques: list[tuple[str, list[str]]] = []
    actual: list[str] | None = None

    for linea in lineas:
        if linea.startswith("## "):
            titulo = linea[3:].strip().lower()
            actual = [linea]
            bloques.append((titulo, actual))
        elif actual is not None:
            actual.append(linea)
        else:
            cabecera.append(linea)

    def clave(item):
        titulo = item[0]
        for i, patron in enumerate(ORDEN_SECCIONES):
            if patron in titulo:
                return i
        return len(ORDEN_SECCIONES)

    ordenados = sorted(bloques, key=clave)
    salida = list(cabecera)
    for _, bloque in ordenados:
        salida.extend(bloque)
    return "\n".join(salida)


def _quitar_titulo_de_la_nota(md_texto: str) -> str:
    """
    Quita el encabezado de nivel 1 con que empieza la nota.

    Es el nombre del archivo de Obsidian ("Aprendizaje - Clase 2026-08-03 -
    ..."), y dentro del documento queda como un titulo que no dice nada y que
    ademas repite al que acaba de poner el generador. Dos encabezados seguidos
    diciendo lo mismo gastan atencion sin dar informacion.
    """
    # El frontmatter va primero: sin quitarlo, la busqueda del titulo se topa
    # con el "---" de apertura y se rinde antes de llegar al encabezado.
    md_texto = _quitar_frontmatter(md_texto)
    lineas = md_texto.splitlines()
    for i, linea in enumerate(lineas):
        if linea.startswith("# "):
            return "\n".join(lineas[:i] + lineas[i + 1:])
        if linea.strip() and not linea.startswith("#"):
            break
    return md_texto


# Lineas que abren un bloque y ademas pueden seguir en los renglones de abajo:
# una viñeta larga, un punto numerado o una cita se cortan igual que un parrafo.
_ABRE_Y_CONTINUA = re.compile(r"^\s*(?:[-*]\s+|\d+\.\s+|>)")

# Lineas que valen por si solas y nunca absorben lo que viene despues.
_LINEA_SUELTA = re.compile(r"^\s*(?:#|\||\$\$|---\s*$)")


def _unir_lineas_de_parrafo(md_texto: str) -> str:
    """
    Une los renglones de un mismo bloque en una sola linea.

    Las notas vienen con el texto cortado cada ochenta o noventa caracteres,
    que es lo normal en markdown y se ve bien en Obsidian. Aqui cada linea se
    convertia en un parrafo suelto de Word, con su espacio debajo, asi que un
    parrafo de seis renglones salia como seis trozos cortados a mitad de frase.
    En markdown un salto simple es continuacion: el bloque termina con una
    linea en blanco.

    Esto vale para parrafos **y tambien para viñetas, puntos numerados y
    citas**. La primera version de esta funcion solo unia la prosa, y como una
    nota real trae decenas de viñetas cortadas (37 en la primera que se midio),
    la segunda mitad de cada una se desprendia igual como parrafo aparte. Se
    veia exactamente igual de mal que el problema que se estaba arreglando.

    Dentro de un bloque cercado (```) no se toca nada: ahi los saltos son parte
    del contenido.
    """
    salida: list[str] = []
    bloque: list[str] = []
    dentro_de_cerca = False

    def volcar():
        if bloque:
            salida.append(" ".join(bloque))
            bloque.clear()

    for linea in md_texto.splitlines():
        desnuda = linea.strip()

        if desnuda.startswith("```"):
            volcar()
            dentro_de_cerca = not dentro_de_cerca
            salida.append(linea)
            continue
        if dentro_de_cerca:
            salida.append(linea)
            continue

        if not desnuda:
            volcar()
            salida.append("")
            continue

        if _LINEA_SUELTA.match(linea):
            volcar()
            salida.append(linea)
            continue

        if _ABRE_Y_CONTINUA.match(linea):
            # Una cita que sigue con otra linea de ">" es la misma cita, no una
            # nueva: se le quita el ">" y se pega a la anterior.
            if bloque and bloque[0].lstrip().startswith(">") and desnuda.startswith(">"):
                bloque.append(desnuda.lstrip(">").strip())
            else:
                volcar()
                bloque.append(linea.rstrip())
            continue

        # Renglon de continuacion: se pega a lo que se venia acumulando, sea un
        # parrafo, una viñeta, un punto numerado o una cita.
        bloque.append(desnuda)

    volcar()
    return "\n".join(salida)


def _quitar_frontmatter(md_texto: str) -> str:
    if md_texto.startswith("---"):
        fin = md_texto.find("\n---", 3)
        if fin != -1:
            return md_texto[fin + 4 :].lstrip("\n")
    return md_texto


def _nivel_encabezado(linea: str) -> int | None:
    if linea.startswith("### "):
        return 3
    if linea.startswith("## "):
        return 2
    if linea.startswith("# "):
        return 1
    return None


def _es_fila_tabla(linea: str) -> bool:
    linea = linea.strip()
    return linea.startswith("|") and linea.endswith("|") and len(linea) > 1


def _es_separador_tabla(linea: str) -> bool:
    return bool(re.fullmatch(r"\|[\s\-:|]+\|", linea.strip()))


def _celdas_de_fila(linea: str) -> list[str]:
    interior = linea.strip().strip("|")
    return [c.strip() for c in interior.split("|")]


def _alinear_izquierda(celda) -> None:
    """
    Las celdas no se justifican. En una columna angosta, justificar reparte el
    sobrante entre dos o tres palabras y deja huecos enormes. El cuerpo del
    documento si va justificado, pero ahi la linea es larga y el sobrante se
    reparte sin que se note.
    """
    for p in celda.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _escribir_celda(celda, texto: str) -> None:
    """
    Una celda que es toda una formula se dibuja; el resto va como texto con
    negritas y subindices.

    Antes esto era `celda.text = texto`, o sea texto plano sin ningun formato.
    La tabla de "que formula uso en cada caso" es justo donde mas formulas hay,
    y era el unico lugar del documento donde se veian crudas, con el guion bajo
    y los parentesis a la vista.
    """
    limpio = texto.strip()
    if formulas.parece_solo_formula(limpio):
        crudo = limpio[1:-1] if limpio.startswith("$") and limpio.endswith("$") else limpio
        if formulas.agregar_formula_en_celda(celda, crudo):
            return
    _agregar_texto_con_negritas(celda.paragraphs[0], texto)


def _agregar_tabla_markdown(doc: Document, filas: list[list[str]]) -> None:
    if not filas:
        return
    encabezado, *resto = filas
    tabla = doc.add_table(rows=1, cols=len(encabezado))
    tabla.style = "Light Grid Accent 1"
    for celda, texto in zip(tabla.rows[0].cells, encabezado):
        celda.text = texto
        _sombrear_celda(celda, COLOR_TABLA_HEADER)
        _alinear_izquierda(celda)
        for p in celda.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
    for fila in resto:
        celdas = tabla.add_row().cells
        for celda, texto in zip(celdas, fila):
            _escribir_celda(celda, texto)
            _alinear_izquierda(celda)


def _leer_bloque_cercado(lineas: list[str], i: int) -> tuple[str, int]:
    """Devuelve el contenido de un bloque ``` y la linea siguiente al cierre."""
    i += 1
    dentro = []
    while i < len(lineas) and not lineas[i].strip().startswith("```"):
        dentro.append(lineas[i])
        i += 1
    return "\n".join(dentro), i + 1


def _agregar_mapa(doc: Document, crudo: str) -> None:
    """Dibuja el mapa conceptual. Si los datos vienen mal, no se pone nada: un
    mapa es un extra y no puede costarle el documento a una clase."""
    import json

    try:
        datos = json.loads(crudo)
        ruta = mapa_visual.dibujar(datos.get("centro", ""), datos.get("ramas") or [])
    except Exception:
        ruta = None
    if ruta is None:
        return

    # La nota suele traer su propio encabezado ("Mapa visual") justo antes del
    # bloque. Poner otro encima deja dos titulos seguidos diciendo lo mismo.
    ultimo = doc.paragraphs[-1] if doc.paragraphs else None
    ya_tiene_titulo = (
        ultimo is not None
        and ultimo.style.name.startswith("Heading")
        and "mapa" in ultimo.text.lower()
    )
    if not ya_tiene_titulo:
        # Nivel 2, no 1: el mapa vive dentro del cuerpo de la clase, y un
        # encabezado de nivel 1 aqui lo sacaria de la seccion a la que pertenece.
        h = doc.add_heading("Mapa de la clase", level=2)
        for run in h.runs:
            run.font.color.rgb = COLOR_ACENTO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ruta), width=Inches(6.3))


def _agregar_markdown(
    doc: Document,
    md_texto: str,
    saltar_secciones: list[str] | None = None,
    solo_secciones: list[str] | None = None,
) -> None:
    """
    Vuelca markdown al documento.

    `saltar_secciones` quita las secciones nombradas. `solo_secciones` hace lo
    contrario y deja pasar unicamente esas, descartando todo lo demas, incluido
    lo que venga antes del primer encabezado. Se usa para la nota de fuente, de
    la que solo interesan tres secciones (ver SECCIONES_FUENTE_EN_DOCX).
    """
    md_texto = _unir_lineas_de_parrafo(_quitar_frontmatter(md_texto))
    saltar_secciones = [s.lower() for s in (saltar_secciones or [])]
    solo_secciones = [s.lower() for s in (solo_secciones or [])]
    nivel_saltando: int | None = None
    # En modo "solo", se arranca descartando hasta encontrar una seccion pedida.
    nivel_dejando_pasar: int | None = None

    lineas = md_texto.splitlines()
    i = 0
    while i < len(lineas):
        linea = lineas[i].rstrip()

        nivel = _nivel_encabezado(linea)
        if nivel is not None:
            titulo_encabezado = linea.lstrip("#").strip()
            if nivel_saltando is not None:
                if nivel <= nivel_saltando:
                    nivel_saltando = None
                else:
                    i += 1
                    continue
            if any(s in titulo_encabezado.lower() for s in saltar_secciones):
                nivel_saltando = nivel
                i += 1
                continue
            if solo_secciones:
                if any(s in titulo_encabezado.lower() for s in solo_secciones):
                    nivel_dejando_pasar = nivel
                elif nivel_dejando_pasar is not None and nivel <= nivel_dejando_pasar:
                    nivel_dejando_pasar = None

        if nivel_saltando is not None:
            i += 1
            continue
        if solo_secciones and nivel_dejando_pasar is None:
            i += 1
            continue

        if not linea.strip():
            i += 1
            continue
        if linea.strip() == "---":
            i += 1
            continue

        if linea.strip().startswith("```mapa"):
            # El modelo entrega la estructura del mapa como datos; el dibujo lo
            # hace mapa_visual.py. Describir un mapa no es un mapa.
            crudo, i = _leer_bloque_cercado(lineas, i)
            _agregar_mapa(doc, crudo)
            continue

        if formulas.es_formula_destacada(linea):
            formulas.agregar_formula_destacada(doc, formulas.texto_de_formula_destacada(linea))
            i += 1
            continue

        if _es_fila_tabla(linea):
            filas = []
            while i < len(lineas) and _es_fila_tabla(lineas[i]):
                if not _es_separador_tabla(lineas[i]):
                    filas.append(_celdas_de_fila(lineas[i]))
                i += 1
            _agregar_tabla_markdown(doc, filas)
            continue

        if linea.startswith("### "):
            h = doc.add_heading(_quitar_numero(linea[4:]), level=3)
        elif linea.startswith("## "):
            h = doc.add_heading(_quitar_numero(linea[3:]), level=2)
            for run in h.runs:
                run.font.color.rgb = COLOR_ACENTO
        elif linea.startswith("# "):
            h = doc.add_heading(linea[2:], level=1)
            for run in h.runs:
                run.font.color.rgb = COLOR_ACENTO
        elif linea.strip().startswith("> [!"):
            tipo = re.match(r"^\s*>\s*\[!([a-zA-Z]+)\]", linea).group(1).lower()
            texto = re.sub(r"^\s*>\s*\[![a-zA-Z]+\]\s*", "", linea)
            # Los dos canales de enfasis del documento. Un callout de otro tipo
            # (la skill no deberia escribirlos, pero las notas viejas los traen)
            # cae en el de verificar, que es el mas discreto: es preferible que
            # un aviso pase desapercibido a que compita con lo que entra en la
            # prueba.
            if tipo == "examen":
                _bloque_destacado(doc, "ENTRA EN LA PRUEBA", texto,
                                  COLOR_EXAMEN_FONDO, "BF8F00")
            else:
                p = doc.add_paragraph()
                _sombrear_parrafo(p, COLOR_VERIFICAR_FONDO)
                _barra_izquierda(p, "808080")
                run = p.add_run("Verificar  ")
                run.bold = True
                run.italic = True
                run_texto = p.add_run(texto)
                run_texto.italic = True
        elif linea.strip().startswith(">"):
            texto = re.sub(r"^\s*>\s*", "", linea)
            p = doc.add_paragraph(style="Intense Quote")
            _agregar_texto_con_negritas(p, texto)
        elif re.match(r"^\s*[-*]\s+", linea):
            texto = re.sub(r"^\s*[-*]\s+", "", linea)
            p = doc.add_paragraph(style="List Bullet")
            _agregar_texto_con_negritas(p, texto)
        elif re.match(r"^\s*\d+\.\s+", linea):
            texto = re.sub(r"^\s*\d+\.\s+", "", linea)
            p = doc.add_paragraph(style="List Number")
            _agregar_texto_con_negritas(p, texto)
        else:
            p = doc.add_paragraph()
            _agregar_texto_con_negritas(p, linea)

        i += 1


def _agregar_llamados(doc: Document, llamados: dict | None) -> None:
    """
    Primera seccion del documento: lo que el profesor pidio.

    Va primero porque es lo unico del material que vence. El resto se puede
    leer cuando sea; una fecha de prueba enterrada en la pagina nueve se lee
    tarde.

    Se muestra siempre, aunque este vacia. Una seccion que desaparece deja al
    estudiante sin saber si el profesor no dijo nada o si el sistema no lo
    detecto, y esas dos cosas no se parecen en nada.

    Cada punto trae la frase textual del profesor. Es la seccion del documento
    con mas riesgo: es la que mas se cree y la que decide que se estudia, asi
    que si no se puede citar, no entra (ver las reglas de honestidad del
    SKILL.md).
    """
    h = doc.add_heading("Lo que el profesor pidió", level=1)
    for run in h.runs:
        run.font.color.rgb = COLOR_ACENTO

    llamados = llamados or {}
    avisos = [a for a in (llamados.get("avisos") or []) if a.get("que")]
    evaluacion = [e for e in (llamados.get("evaluacion") or []) if e.get("tema")]

    if not avisos and not evaluacion:
        p = doc.add_paragraph()
        run = p.add_run(
            "En esta clase el profesor no anunció fechas, entregas ni contenidos "
            "de evaluación."
        )
        run.italic = True
        return

    for aviso in avisos:
        cuando = str(aviso.get("cuando", "")).strip()
        texto = str(aviso.get("que", "")).strip()
        if cuando:
            texto = f"{texto} ({cuando})"
        p = _bloque_destacado(doc, "", texto, COLOR_EXAMEN_FONDO, "BF8F00")
        _agregar_textual(doc, aviso)

    if evaluacion:
        h2 = doc.add_heading("Lo que dijo que entra en evaluación", level=2)
        for run in h2.runs:
            run.font.color.rgb = COLOR_ACENTO
        intro = doc.add_paragraph()
        run = intro.add_run(
            "Si el profesor se tomó el trabajo de decirlo en voz alta, es lo "
            "primero que hay que estudiar."
        )
        run.italic = True

        for item in evaluacion:
            p = _bloque_destacado(
                doc, "ENTRA", str(item.get("tema", "")).strip(),
                COLOR_EXAMEN_FONDO, "BF8F00",
            )
            _agregar_textual(doc, item)


def _agregar_textual(doc: Document, item: dict) -> None:
    """La cita del profesor debajo de cada punto, y el aviso cuando no se pudo
    escuchar bien. Sin la cita el punto no se puede verificar, que es justo lo
    que hace confiable a esta seccion."""
    textual = str(item.get("textual", "")).strip()
    if textual:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(f'"{textual}"')
        run.italic = True
        run.font.size = Pt(10)
    if item.get("seguro") is False:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        run = p.add_run(
            "El audio no permite estar seguro de este punto. Confírmalo antes "
            "de organizar el estudio con él."
        )
        run.italic = True
        run.font.size = Pt(10)


def _agregar_contexto_previo(doc: Document, texto: str) -> None:
    """
    Va al principio del documento, antes que nada del contenido de la clase.

    Existe porque una clase sobre algo nuevo se escucha sin red: el profesor da
    por sabido lo que vio en cursos anteriores, y el estudiante que no lo tiene
    fresco pierde la primera media hora tratando de ubicarse. Esta seccion se
    lee primero y da justo el piso necesario para que el resto se entienda.

    Es lo unico del documento que NO sale de la transcripcion, asi que se
    marca como tal: el resto del material vale porque es lo que el profesor
    dijo, y esa distincion no se puede difuminar.
    """
    h = doc.add_heading("Antes de empezar: lo que conviene tener claro", level=1)
    for run in h.runs:
        run.font.color.rgb = COLOR_ACENTO

    aviso = doc.add_paragraph()
    _sombrear_parrafo(aviso, COLOR_CONTEXTO_FONDO)
    run = aviso.add_run(
        "Esta seccion no es parte de la clase: es contexto agregado para que lo "
        "que viene se entienda desde cero. Leela primero."
    )
    run.italic = True

    _agregar_markdown(doc, _quitar_titulo_de_la_nota(texto))


def _agregar_conceptos_repetidos(doc: Document, conceptos: list[dict]) -> None:
    h = doc.add_heading("Conceptos mas repetidos por el profesor", level=1)
    for run in h.runs:
        run.font.color.rgb = COLOR_ACENTO
    doc.add_paragraph(
        "Estos conceptos volvieron una y otra vez durante la clase: son candidatos "
        "fuertes para la prueba."
    ).italic = True

    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Light Grid Accent 1"
    encabezados = tabla.rows[0].cells
    encabezados[0].text = "Concepto"
    encabezados[1].text = "Por que se repite"
    for celda in encabezados:
        _sombrear_celda(celda, COLOR_TABLA_HEADER)
        _alinear_izquierda(celda)
        for p in celda.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True

    for item in conceptos:
        fila = tabla.add_row().cells
        fila[0].text = str(item.get("concepto", ""))
        fila[1].text = str(item.get("por_que", ""))
        for celda in fila:
            _alinear_izquierda(celda)


def _aplicar_tipografia(doc: Document) -> None:
    """
    Margenes, interlineado y cuerpo, segun references/diseno-documento.md.

    Lo que mas cambia la lectura no es la fuente, es el ancho de la linea: por
    encima de unos 75 caracteres el ojo pierde el renglon al volver a la
    izquierda. Word por defecto deja unos 95.
    """
    for seccion in doc.sections:
        seccion.left_margin = Cm(MARGEN_LATERAL_CM)
        seccion.right_margin = Cm(MARGEN_LATERAL_CM)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = INTERLINEADO
    normal.paragraph_format.space_after = Pt(10)
    if not JUSTIFICADO:
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return

    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Particion de palabras al justificar. Sin esto, Word estira los espacios
    # de la linea hasta el margen y aparecen los "rios" blancos que cruzan el
    # parrafo de arriba abajo, que es justo lo que hace incomodo el texto
    # justificado. Con particion, el sobrante se reparte en la palabra.
    settings = doc.settings.element
    for nombre, valor in (("autoHyphenation", "1"), ("hyphenationZone", "284"),
                          ("consecutiveHyphenLimit", "2")):
        el = OxmlElement(f"w:{nombre}")
        el.set(qn("w:val"), valor)
        settings.append(el)


def generar_docx(
    trabajo: dict,
    titulo: str,
    texto_fuente: str,
    texto_aprendizaje: str,
    conceptos_repetidos: list[dict],
    config: dict,
    texto_contexto: str = "",
    llamados: dict | None = None,
    ruta_destino: Path | None = None,
) -> Path:
    """
    El orden de las secciones esta fijado en references/diseno-documento.md y
    cada una hace un trabajo distinto. `ruta_destino` permite escribir en otro
    lado sin pisar el documento de la clase, que es lo que usa la regeneracion
    de clases ya procesadas.
    """
    doc = Document()
    _aplicar_tipografia(doc)

    portada = doc.add_heading(f"Clase {trabajo['numero_clase']:02d} - {trabajo['ramo']}", level=0)
    portada.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo = doc.add_paragraph(f"{titulo} ({trabajo['fecha']})")
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.runs[0].italic = True
    doc.add_page_break()

    # Primero lo unico que vence.
    _agregar_llamados(doc, llamados)
    doc.add_page_break()

    if texto_contexto:
        _agregar_contexto_previo(doc, texto_contexto)
        doc.add_page_break()

    if conceptos_repetidos:
        _agregar_conceptos_repetidos(doc, conceptos_repetidos)

    if texto_aprendizaje:
        # Encabezado propio para el cuerpo. Sin el, las secciones de la nota
        # quedaban colgando del titulo anterior ("Conceptos mas repetidos"), que
        # no es su padre: la materia parecia una subseccion de la tabla de
        # conceptos.
        doc.add_page_break()
        h = doc.add_heading("Material de estudio", level=1)
        for run in h.runs:
            run.font.color.rgb = COLOR_ACENTO
        cuerpo = _reordenar_secciones(_quitar_titulo_de_la_nota(texto_aprendizaje))
        _agregar_markdown(
            doc, cuerpo, saltar_secciones=_secciones_a_omitir(texto_aprendizaje)
        )

    if texto_fuente:
        doc.add_page_break()
        h = doc.add_heading("Respaldo de la clase", level=1)
        for run in h.runs:
            run.font.color.rgb = COLOR_ACENTO
        intro = doc.add_paragraph()
        run = intro.add_run(
            "Las definiciones tal como las dijo el profesor, lo que hubo que "
            "reconstruir, y lo que la grabación no permitió recuperar."
        )
        run.italic = True
        _agregar_markdown(
            doc, _quitar_titulo_de_la_nota(texto_fuente),
            solo_secciones=SECCIONES_FUENTE_EN_DOCX,
        )

    if ruta_destino is not None:
        ruta = Path(ruta_destino)
        ruta.parent.mkdir(parents=True, exist_ok=True)
    else:
        output_dir = Path(config["rutas"]["output"]) / trabajo["ramo"]
        output_dir.mkdir(parents=True, exist_ok=True)
        base = nombre_base(trabajo["numero_clase"], trabajo["fecha"], titulo)
        ruta = output_dir / f"{base}.docx"
    doc.save(str(ruta))
    # Las imagenes ya quedaron incrustadas en el .docx.
    formulas.limpiar_temporales()
    mapa_visual.limpiar_temporales()
    return ruta
